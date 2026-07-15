"""

FastAPI 路由定义

提供 REST API 接口供外部调用

包含：认证（JWT）、聊天（含流式）、文档管理、会话管理、模型管理、统计



优化:

- [#20] 可观测性：请求日志中间件 + 性能指标

- [#22] 配置中心：运行时热更新配置 API

- [#23] API 分页：对话列表/文档列表支持分页

- [#24] 健康检查增强：检查 ChromaDB/LLM API/磁盘等依赖

"""

import os

import asyncio

import time

import shutil

import json

import base64

import logging

from typing import Optional



from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends, Request, Query, Header

from fastapi.responses import StreamingResponse, Response, FileResponse

from pydantic import BaseModel

from urllib.parse import unquote



from app.agent.core import chat, chat_stream_generator, chat_stream_generator_multimodal, reset_agent


# [BUG FIX v5] 可复用的 SSE 流式包装器：客户端断开时真正取消 Agent 执行
async def _sse_stream_wrapper(generator_factory, request: Request, session_id: str, start_time: float, endpoint: str = "/chat/stream"):
    """将 chat_stream_generator 包装为 Queue+Producer Task 模式
    
    当客户端断开时，cancel producer_task 可真正终止 Agent 执行。
    generator_factory: 无参数的 async generator 工厂函数，如 lambda: chat_stream_generator(...)
    
    [v6 优化] 去掉内层 create_task + sleep 轮询，改用 asyncio.wait_for(queue.get, timeout=0.5)
    避免每个 SSE 连接每次消费创建 throwaway task + 50ms 忙等待
    """
    queue = asyncio.Queue()
    stream_done = object()
    cancelled_by_client = False
    
    async def produce():
        nonlocal cancelled_by_client
        try:
            async for chunk in generator_factory():
                if await request.is_disconnected():
                    cancelled_by_client = True
                    logger.info(f"SSE客户端断开，正在终止Agent执行: session={session_id}")
                    break
                await queue.put(chunk)
            await queue.put(stream_done)
        except asyncio.CancelledError:
            logger.info(f"Agent执行任务被取消: session={session_id}")
            raise
        except Exception as e:
            logger.exception(f"SSE生产者异常: session={session_id}")
            await queue.put({'type': 'error', 'content': str(e)})
            await queue.put(stream_done)
    
    producer_task = asyncio.create_task(produce())
    
    try:
        while True:
            # [v6 优化] 直接 await queue.get() + 超时检测断开，不创建中间 task
            if await request.is_disconnected():
                cancelled_by_client = True
                logger.info(f"SSE客户端断开，正在取消Agent执行: session={session_id}")
                producer_task.cancel()
                try:
                    await asyncio.shield(producer_task)
                except asyncio.CancelledError:
                    pass
                return
            
            try:
                chunk = await asyncio.wait_for(queue.get(), timeout=0.5)
            except asyncio.TimeoutError:
                # 超时后回到循环顶部检查 disconnect
                continue
            
            if chunk is stream_done:
                break
            if isinstance(chunk, dict) and chunk.get('type') == 'error':
                yield f"data: {json.dumps(chunk, ensure_ascii=False)}\n\n"
                yield f"data: {json.dumps({'type': 'done'}, ensure_ascii=False)}\n\n"
                break
            yield f"data: {json.dumps(chunk, ensure_ascii=False)}\n\n"
    except asyncio.CancelledError:
        logger.info(f"SSE流被取消（外部信号）: session={session_id}")
        producer_task.cancel()
        try:
            await asyncio.shield(producer_task)
        except asyncio.CancelledError:
            pass
        return
    finally:
        if not producer_task.done():
            producer_task.cancel()
    
    # 更新会话时间
    try:
        parts = session_id.split("_", 1)
        if len(parts) == 2:
            update_chat_time(parts[0], session_id)
    except Exception:
        pass

    # [BUG FIX] 流结束后立即flush，确保对话数据写入磁盘
    # 避免多worker进程下导出时读到过时数据
    try:
        flush_session(session_id)
    except Exception:
        pass

    _record_request(endpoint, time.time() - start_time)
    if cancelled_by_client:
        logger.info(f"SSE流完成（客户端主动断开）: session={session_id}")

from app.rag.document import index_document, search_documents, list_indexed_documents, delete_document, update_document, delete_agent_collection, list_all_collections, load_document, export_document_as_docx, reindex_all_documents, get_indexing_mode, _get_export_dir, cleanup_export_files, _load_keyword_index, get_vector_store, _get_agent_dir

from app.auth.user_manager import login_user, register_user, get_user_role, is_admin, list_all_users, delete_user, update_user_role, reset_user_password

from app.auth.jwt_handler import create_token, verify_token, get_username_from_token, get_role_from_token

from app.memory.manager import (

    get_history_messages, get_history_messages_from_file, clear_session_history,
    flush_session,

    create_chat, list_chats, delete_chat, rename_chat, update_chat_time,

)

from app.config import settings, AVAILABLE_MODELS, get_current_model, set_current_model

from app.utils.stats import record_message, record_session, get_stats

from app.agent.storage import sync_agents as storage_sync_agents, load_agents as storage_load_agents



logger = logging.getLogger(__name__)



# 文件大小限制：50MB

MAX_FILE_SIZE = 50 * 1024 * 1024



router = APIRouter()





# ===== [#20] 可观测性：请求计时 + 性能日志 =====

import threading



_request_stats = {

    "total_requests": 0,

    "total_errors": 0,

    "avg_response_time": 0.0,

    "endpoint_stats": {},  # path -> {count, avg_time, errors}

}

_request_stats_lock = threading.Lock()



# [性能修复] 端点统计上限，避免长时间运行后内存无限增长

_MAX_ENDPOINT_STATS = 50





def _record_request(path: str, duration: float, is_error: bool = False):

    """记录请求统计（线程安全）"""

    with _request_stats_lock:

        _request_stats["total_requests"] += 1

        if is_error:

            _request_stats["total_errors"] += 1

        

        # 更新平均响应时间

        total = _request_stats["total_requests"]

        prev_avg = _request_stats["avg_response_time"]

        _request_stats["avg_response_time"] = prev_avg + (duration - prev_avg) / total

        

        # 端点统计

        if path not in _request_stats["endpoint_stats"]:

            # [性能修复] 超过上限时淘汰请求量最少的端点

            if len(_request_stats["endpoint_stats"]) >= _MAX_ENDPOINT_STATS:

                min_path = min(_request_stats["endpoint_stats"], 

                              key=lambda k: _request_stats["endpoint_stats"][k]["count"])

                del _request_stats["endpoint_stats"][min_path]

            _request_stats["endpoint_stats"][path] = {"count": 0, "avg_time": 0.0, "errors": 0}

        ep = _request_stats["endpoint_stats"][path]

        ep["count"] += 1

        prev = ep["avg_time"]

        ep["avg_time"] = prev + (duration - prev) / ep["count"]

        if is_error:

            ep["errors"] += 1





# ===== JWT 认证依赖 =====

def get_current_user(request: Request) -> str:

    """

    从请求中提取当前用户名（JWT Token 验证）

    不强制认证，但如果有 Token 则验证

    注意：已移除查询参数回退，防止认证绕过

    """

    auth_header = request.headers.get("Authorization", "")

    if auth_header.startswith("Bearer "):

        token = auth_header[7:]

        username = get_username_from_token(token)

        if username:

            return username

    return ""





def require_auth(request: Request) -> str:

    """

    强制要求 JWT 认证

    返回已认证的用户名

    """

    auth_header = request.headers.get("Authorization", "")

    if auth_header.startswith("Bearer "):

        token = auth_header[7:]

        username = get_username_from_token(token)

        if username:

            return username

    raise HTTPException(status_code=401, detail="未认证，请重新登录")









def require_admin(request: Request) -> str:
    """
    强制要求 JWT 认证且为管理员角色
    返回已认证的管理员用户名
    """
    auth_header = request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        token = auth_header[7:]
        username = get_username_from_token(token)
        role = get_role_from_token(token)
        if username and role == "admin":
            return username
        elif username:
            raise HTTPException(status_code=403, detail="权限不足，需要管理员权限")
    raise HTTPException(status_code=401, detail="未认证，请重新登录")


# ===== 请求/响应模型 =====

class ChatRequest(BaseModel):

    """聊天请求"""

    message: str

    session_id: str = "default"

    web_search: bool = False

    mode: str = "agent"  # agent / chat

    deep_think: bool = False

    agent_id: str = None  # 智能体ID，用于知识库隔离

    agent_task: str = None  # 智能体任务描述，用于动态系统提示词
    skill: str = None  # [方案B] 前端选择的技能ID（如 8d-skill），用于注入 SKILL.md + 模板





class ChatResponse(BaseModel):

    """聊天响应"""

    response: str

    session_id: str





class SearchRequest(BaseModel):

    """文档搜索请求"""

    query: str

    top_k: int = 3





class LoginRequest(BaseModel):

    """登录请求"""

    username: str

    password: str





class RegisterRequest(BaseModel):

    """注册请求"""

    username: str

    password: str





class ModelSetRequest(BaseModel):

    """设置模型请求"""

    model_id: str





class RenameRequest(BaseModel):

    """重命名会话请求"""

    username: str

    chat_id: str

    new_title: str





# [#22] 配置中心请求模型

class ConfigUpdateRequest(BaseModel):

    """配置更新请求"""

    key: str  # 配置项名称，如 LLM_MODEL, MAX_TOOL_ROUNDS 等

    value: str  # 新值（字符串形式，内部转换）





class ModifyDocumentRequest(BaseModel):

    """修改知识库文档请求"""

    content: str  # 新的文档内容（纯文本）

    append: bool = False  # 是否追加内容（True=在原文末尾追加，False=替换全部内容）

    return_docx: bool = False  # 是否同时返回修改后的docx文件下载链接

    agent_id: str = None  # 智能体ID，用于知识库隔离

    session_id: str = ""  # [Bug 修复] 会话ID，用于导出文件按会话隔离目录





class ExportDocumentRequest(BaseModel):

    """导出/生成文档请求"""

    content: str  # 文档内容（纯文本）

    filename: str = ""  # 输出文件名（含扩展名），为空则自动生成

    title: str = ""  # 文档标题，为空则使用filename

    session_id: str = ""  # [Bug 修复] 会话ID，用于导出文件按会话隔离目录





# ===== 认证接口 =====



@router.post("/auth/login", summary="用户登录")

async def auth_login(req: LoginRequest):

    """用户登录验证，返回 JWT Token（含角色信息）"""

    start = time.time()

    try:

        result = login_user(req.username, req.password)

        if result.get("success"):

            # 签发 JWT Token（包含角色信息）
            user_role = result.get("role", "user")

            token = create_token(req.username, role=user_role)

            result["token"] = token

        return result

    finally:

        _record_request("/auth/login", time.time() - start)





@router.post("/auth/register", summary="用户注册（已禁用，仅管理员可通过 /admin/users 创建）")

async def auth_register(req: RegisterRequest):

    """用户注册已禁用，前端不提供注册入口，新用户只能由管理员在后端创建"""

    raise HTTPException(status_code=403, detail="注册功能已禁用，请联系管理员创建账号")





@router.get("/auth/me", summary="验证 Token 有效性")

async def auth_me(request: Request):

    """验证当前 JWT Token 是否有效"""

    try:

        username = require_auth(request)

        return {"valid": True, "username": username}

    except HTTPException:

        return {"valid": False, "username": None}








# ===== 管理员用户管理接口 =====



class AdminCreateUserRequest(BaseModel):

    """管理员创建用户请求"""

    username: str

    password: str

    role: str = "user"  # admin 或 user，默认 user



class AdminUpdateRoleRequest(BaseModel):

    """管理员修改用户角色请求"""

    role: str  # admin 或 user



class AdminResetPasswordRequest(BaseModel):

    """管理员重置用户密码请求"""

    new_password: str



@router.post("/admin/users", summary="管理员创建新用户")

async def admin_create_user(req: AdminCreateUserRequest, admin: str = Depends(require_admin)):

    """
    管理员创建新用户（仅管理员可用）
    前端不提供注册入口，所有新用户必须由管理员通过此接口创建
    """

    start = time.time()

    try:

        result = register_user(req.username, req.password, role=req.role)

        if result.get("success"):

            logger.info(f"管理员 {admin} 创建了新用户: {req.username}, 角色: {req.role}")

        return result

    finally:

        _record_request("/admin/users", time.time() - start)



@router.get("/admin/users", summary="管理员获取用户列表（含明文密码）")

async def admin_list_users(admin: str = Depends(require_admin)):

    """
    管理员获取所有用户信息（含明文密码），仅管理员可用
    """

    start = time.time()

    try:

        users = list_all_users()

        return {"success": True, "users": users, "total": len(users)}

    finally:

        _record_request("/admin/users", time.time() - start)



@router.delete("/admin/users/{username}", summary="管理员删除用户")

async def admin_delete_user(username: str, admin: str = Depends(require_admin)):

    """
    管理员删除指定用户（不允许删除 admin 账号），仅管理员可用
    """

    start = time.time()

    try:

        result = delete_user(username)

        if result.get("success"):

            logger.info(f"管理员 {admin} 删除了用户: {username}")

        return result

    finally:

        _record_request("/admin/users/delete", time.time() - start)



@router.put("/admin/users/{username}/role", summary="管理员修改用户角色")

async def admin_update_user_role(username: str, req: AdminUpdateRoleRequest, admin: str = Depends(require_admin)):

    """
    管理员修改用户角色（admin/user），仅管理员可用
    """

    start = time.time()

    try:

        result = update_user_role(username, req.role)

        if result.get("success"):

            logger.info(f"管理员 {admin} 修改用户 {username} 角色为: {req.role}")

        return result

    finally:

        _record_request("/admin/users/role", time.time() - start)



@router.put("/admin/users/{username}/password", summary="管理员重置用户密码")

async def admin_reset_user_password(username: str, req: AdminResetPasswordRequest, admin: str = Depends(require_admin)):

    """
    管理员重置指定用户密码，仅管理员可用
    """

    start = time.time()

    try:

        result = reset_user_password(username, req.new_password)

        if result.get("success"):

            logger.info(f"管理员 {admin} 重置了用户 {username} 的密码")

        return result

    finally:

        _record_request("/admin/users/password", time.time() - start)

# ===== 聊天接口 =====



@router.post("/chat", response_model=ChatResponse, summary="与 Agent 对话（非流式）")

async def chat_api(req: ChatRequest, username: str = Depends(get_current_user)):

    """

    核心接口：与文档助手 Agent 对话（非流式）



    - 支持 RAG 文档问答

    - 支持员工信息查询

    - 支持多轮对话

    """

    start = time.time()

    is_error = False

    try:

        # [BUG FIX v6] chat() 是同步阻塞函数（内部 llm.invoke / agent.invoke），

        # 必须用 asyncio.to_thread 放到线程池，否则阻塞整个事件循环导致所有请求卡死

        response = await asyncio.to_thread(chat, req.message, req.session_id, web_search=req.web_search, mode=req.mode, deep_think=req.deep_think, agent_id=_user_agent_id(req.agent_id, username), agent_task=req.agent_task, skill=req.skill)

        # 更新会话时间

        try:

            parts = req.session_id.split("_", 1)

            if len(parts) == 2:

                update_chat_time(parts[0], req.session_id)

        except Exception:

            pass

        # 记录统计

        record_message(username=username or "anonymous", model_id=get_current_model())

        return ChatResponse(response=response, session_id=req.session_id)

    except Exception as e:

        is_error = True

        raise HTTPException(status_code=500, detail=f"Agent 处理失败: {str(e)}")

    finally:

        _record_request("/chat", time.time() - start, is_error=is_error)





@router.post("/chat/stream", summary="与 Agent 对话（流式 SSE）")

async def chat_stream_api(req: ChatRequest, request: Request, username: str = Depends(get_current_user)):

    """

    流式对话接口：逐 token 输出，同时显示工具调用进度

    返回 Server-Sent Events (SSE) 流

    

    性能优化：检测客户端断开，避免服务端空转消耗资源

    

    BUG FIX v5：客户端断开时通过取消 asyncio.Task 真正终止 Agent 执行，

    不再只是 break 退出循环（旧方式会导致 LangGraph 后台继续调用 LLM，消耗 rate limit）

    """

    start = time.time()

    # 记录统计

    record_message(username=username or "anonymous", model_id=get_current_model())



    generator_factory = lambda: chat_stream_generator(req.message, req.session_id, web_search=req.web_search, mode=req.mode, deep_think=req.deep_think, agent_id=_user_agent_id(req.agent_id, username), agent_task=req.agent_task, skill=req.skill)



    return StreamingResponse(

        _sse_stream_wrapper(generator_factory, request, req.session_id, start, endpoint="/chat/stream"),

        media_type="text/event-stream",

        headers={

            "Cache-Control": "no-cache",

            "Connection": "keep-alive",

            "X-Accel-Buffering": "no",

        },

    )





@router.post("/chat-with-file/stream", summary="带文件的流式对话")

async def chat_with_file_stream(

    request: Request,

    file: UploadFile = File(...),

    message: str = Form(""),

    session_id: str = Form("default"),

    web_search: bool = Form(False),

    mode: str = Form("agent"),

    deep_think: bool = Form(False),

    agent_id: str = Form(None),

    agent_task: str = Form(None),

    skill: str = Form(""),  # [方案B] 前端选择的技能ID（如 8d-skill）

    store_to_kb: str = Form("true"),

    category: str = Form(""),  # [Bug 1 修复] 补全缺失的 category 参数

    username: str = Depends(get_current_user),

):

    """

    带文件的流式对话：支持图片和文档

    - 图片（png/jpg/jpeg/gif/bmp/webp）：转为base64传给LLM分析

    - 文档（pdf/txt/docx）：索引后基于内容回答

    - 其他文件：读取文本内容（如有）传给LLM

    返回 Server-Sent Events (SSE) 流

    """

    start = time.time()

    # 记录统计

    record_message(username=username or "anonymous", model_id=get_current_model())



    ext = os.path.splitext(file.filename)[1].lower()

    image_exts = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}

    doc_exts = {".pdf", ".txt", ".docx", ".xlsx", ".xls"}

    code_exts = {".py", ".js", ".html", ".css", ".json", ".md", ".csv", ".xlsx", ".xls", ".doc", ".ppt", ".pptx"}



    # 文件大小检查

    file_content_raw = await file.read()

    if len(file_content_raw) > MAX_FILE_SIZE:

        raise HTTPException(status_code=413, detail=f"文件大小超过限制（最大 50MB），当前文件: {len(file_content_raw) // 1024 // 1024}MB")

    # 重置文件指针

    await file.seek(0)



    logger.info(f"收到文件上传: {file.filename}, 大小: {len(file_content_raw)} bytes")



    if ext in image_exts:

        # 图片文件：用多模态消息格式传给LLM做视觉分析（复用已读取的 file_content_raw）

        b64 = base64.b64encode(file_content_raw).decode("utf-8")

        mime_map = {

            ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",

            ".gif": "image/gif", ".bmp": "image/bmp", ".webp": "image/webp",

        }

        mime_type = mime_map.get(ext, "image/png")

        # 构建多模态消息内容

        image_url = f"data:{mime_type};base64,{b64}"

        multimodal_content = [

            {"type": "text", "text": f"[用户上传了图片: {file.filename}]\n\n{message or '请描述这张图片'}"},

            {"type": "image_url", "image_url": {"url": image_url}},

        ]

        # 直接调用多模态流式生成

        return StreamingResponse(

            _sse_stream_wrapper(

                lambda: chat_stream_generator_multimodal(multimodal_content, session_id, agent_id=_user_agent_id(agent_id, username), agent_task=agent_task, skill=skill or None),

                request, session_id, start, endpoint="/chat-with-file/stream"

            ),

            media_type="text/event-stream",

            headers={

                "Cache-Control": "no-cache",

                "Connection": "keep-alive",

                "X-Accel-Buffering": "no",

            },

        )



    elif ext in doc_exts:

        # 普通聊天模式（无 agent_id）：不存入知识库，只临时读取文件内容回答

        # 文件保存到临时目录，删除会话时自动清理

        # URL解码文件名：浏览器上传的中文文件名可能是URL编码的，统一解码

        decoded_filename = unquote(file.filename)

        if store_to_kb == "true" and agent_id:

            # 知识库模式 ON + 有 agent_id：文件存到智能体知识库目录（删对话不删除）

            agent_dir = os.path.join(settings.DOCUMENTS_DIR, f"agent_{_user_agent_id(agent_id, username)}")

            os.makedirs(agent_dir, exist_ok=True)

            file_path = os.path.join(agent_dir, decoded_filename)

        else:

            # 其他情况（普通模式 / 知识库模式OFF / 无agent_id）：文件存到临时目录，删除对话时自动清理

            temp_dir = os.path.join(settings.DATA_DIR, "temp", session_id)

            os.makedirs(temp_dir, exist_ok=True)

            file_path = os.path.join(temp_dir, decoded_filename)

        with open(file_path, "wb") as f:

            shutil.copyfileobj(file.file, f)



        if store_to_kb == "true" and agent_id:

            # 知识库模式 ON + 有 agent_id：索引到智能体知识库

            try:

                index_result = await asyncio.to_thread(index_document, file_path, decoded_filename, agent_id=_user_agent_id(agent_id, username), category=category)

                indexing_mode = index_result.get('indexing_mode', 'unknown')

                logger.info(f"文件已索引到知识库: {file.filename}, agent_id={agent_id}, 分块数={index_result.get('chunks', 0)}, 索引模式={indexing_mode}")

            except Exception as e:

                os.remove(file_path)

                raise HTTPException(status_code=500, detail=f"文档索引失败: {str(e)}")

            full_message = f"[用户上传了文档: {file.filename}]\n\n{message}"

        else:

            # 普通模式或 store_to_kb=false：只读取内容回答，不存入知识库

            try:

                docs = await asyncio.to_thread(load_document, file_path)

                text = "\n".join([doc.page_content for doc in docs])

                full_message = f"[用户上传了文档: {file.filename}]\n\n文档内容：\n{text[:8000]}\n\n{message}"

                mode_label = "普通聊天（不存知识库）" if not agent_id else "知识库模式OFF"

                logger.info(f"文件仅读取内容（{mode_label}）: {file.filename}")

            except Exception as e:

                os.remove(file_path)

                raise HTTPException(status_code=500, detail=f"文档读取失败: {str(e)}")



    elif ext in code_exts:

        # 代码/其他文本文件：读取内容传给LLM

        try:

            file_content = await file.read()

            text = file_content.decode("utf-8", errors="replace")

            full_message = f"[用户上传了文件: {file.filename}]\n\n文件内容：\n```\n{text[:8000]}\n```\n\n{message}"

        except Exception:

            full_message = f"[用户上传了文件: {file.filename}，但无法读取内容]\n\n{message}"

    else:

        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")



    # 流式回答

    full_message_local = full_message  # 避免闭包引用问题

    aid_local = agent_id if agent_id else None

    atask_local = agent_task if agent_task else None



    return StreamingResponse(

        _sse_stream_wrapper(

            lambda: chat_stream_generator(full_message_local, session_id, web_search=web_search, mode=mode, deep_think=deep_think, agent_id=_user_agent_id(aid_local, username), agent_task=atask_local, skill=skill or None),

            request, session_id, start, endpoint="/chat-with-file/stream"

        ),

        media_type="text/event-stream",

        headers={

            "Cache-Control": "no-cache",

            "Connection": "keep-alive",

            "X-Accel-Buffering": "no",

        },

    )





# ===== 文档管理接口 =====



@router.post("/upload", summary="上传文档到知识库")

async def upload_document(file: UploadFile = File(...), agent_id: str = Form(None), category: str = Form(''), subcategory: str = Form(''), username: str = Depends(require_auth)):

    """

    上传文档并自动索引到向量数据库（需登录认证）

    支持 PDF、TXT、MD、DOCX、XLSX、XLS、图片(PNG/JPG/JPEG/GIF/BMP/WebP) 格式

    必须指定 agent_id（普通聊天模式无知识库，不支持上传到知识库）

    """

    # 普通聊天模式无知识库，必须指定 agent_id

    if not agent_id:

        raise HTTPException(

            status_code=400,

            detail="请先选择一个智能体再上传文档到知识库。普通聊天模式不支持知识库功能。",

        )

    # [用户隔离] agent_id 加用户名前缀
    agent_id = _user_agent_id(agent_id, username)



    # 检查文件格式

    allowed_ext = {".pdf", ".txt", ".md", ".docx", ".doc", ".xlsx", ".xls", ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}

    ext = os.path.splitext(file.filename)[1].lower()

    if ext not in allowed_ext:

        raise HTTPException(

            status_code=400,

            detail=f"不支持的文件格式: {ext}，仅支持 {allowed_ext}",

        )



    # 文件大小检查

    file_content_raw = await file.read()

    if len(file_content_raw) > MAX_FILE_SIZE:

        raise HTTPException(status_code=413, detail=f"文件大小超过限制（最大 50MB）")

    await file.seek(0)



    logger.info(f"知识库上传文档: {file.filename}, 大小: {len(file_content_raw)} bytes")



    # 保存文件 - 使用per-agent子目录实现文件隔离

    if agent_id:

        agent_dir = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}")

        if category:

            agent_dir = os.path.join(agent_dir, category)

        if subcategory:

            agent_dir = os.path.join(agent_dir, subcategory)

        os.makedirs(agent_dir, exist_ok=True)

        # URL解码文件名：浏览器上传的中文文件名可能是URL编码的，统一解码

        decoded_filename = unquote(file.filename)

        file_path = os.path.join(agent_dir, decoded_filename)

    else:

        # URL解码文件名

        decoded_filename = unquote(file.filename)

        file_path = os.path.join(settings.DOCUMENTS_DIR, decoded_filename)

    with open(file_path, "wb") as f:

        shutil.copyfileobj(file.file, f)



    # 索引文档（[#11] 自动降级：embedding不可用时切换为关键词索引）

    try:

        # [性能修复] 用 asyncio.to_thread 在线程池中执行同步 index_document，

        # 避免文件加载+分块+Embedding API调用阻塞整个事件循环

        result = await asyncio.to_thread(index_document, file_path, decoded_filename, agent_id=agent_id, category=category or None, subcategory=subcategory or None)

        indexing_mode_result = result.get('indexing_mode', 'unknown')

        logger.info(f"文档索引完成: {file.filename}, agent_id={agent_id}, 分块数={result.get('chunks', 0)}, 索引模式={indexing_mode_result}")

        return {"status": "success", "detail": result}

    except Exception as e:

        # 索引失败则删除文件

        os.remove(file_path)

        raise HTTPException(status_code=500, detail=f"文档索引失败: {str(e)}")





@router.post("/search", summary="搜索文档内容")

async def search_api(req: SearchRequest, agent_id: str = Query(None, description="智能体ID，为空时搜全局知识库"), username: str = Depends(require_auth)):

    """在文档库中搜索相关内容（支持按智能体隔离）"""
    agent_id = _user_agent_id(agent_id, username)

    # 普通聊天模式没有知识库

    if not agent_id:

        return {"query": req.query, "results": [], "message": "普通聊天模式没有知识库，请先选择一个智能体"}

    results = search_documents(req.query, req.top_k, agent_id=agent_id)

    return {"query": req.query, "results": results}





@router.get("/documents", summary="列出所有已索引文档")

async def list_documents(

    page: int = Query(1, ge=1, description="页码"),          # [#23] 分页

    page_size: int = Query(20, ge=1, le=100, description="每页数量"),

    agent_id: str = Query(None, description="智能体ID，为空时查全局知识库"),

    category: str = Query(None, description="文件分类（手册/程序文件/三层次文件/记录表格/其他）"),

    subcategory: str = Query(None, description="二级子目录名（必须同时指定 category）"),

    username: str = Depends(require_auth),

):

    """获取知识库中所有文档列表（支持分页，按智能体隔离）



    [#11] 同时扫描关键词索引和磁盘文件，确保关键词模式下也能正确列出文档

    

    注意：普通聊天模式（agent_id=None）没有知识库，返回空列表

    """

    # 普通聊天模式没有知识库

    if not agent_id:

        return {

            "documents": [],

            "count": 0,

            "total": 0,

            "page": page,

            "page_size": page_size,

        }

    # [用户隔离] agent_id 加用户名前缀
    agent_id = _user_agent_id(agent_id, username)

    docs = list_indexed_documents(agent_id=agent_id, category=category, subcategory=subcategory)

    # 额外扫描：list_indexed_documents 已扫描 .pdf/.txt/.docx，

    # 这里补充扫描更多文件类型（代码文件、Office文档等）

    extra_extensions = {'.csv', '.xlsx', '.xls', '.doc', '.ppt', '.pptx', '.md', '.py', '.js', '.html', '.css', '.json'}

    indexed_filenames = set()

    for doc in docs:

        if isinstance(doc, dict) and doc.get('filename'):

            indexed_filenames.add(doc['filename'])

        elif isinstance(doc, str):

            indexed_filenames.add(doc)



    # 扫描对应的目录（仅补充额外格式的文件）

    if agent_id:

        scan_dir = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}")

    else:

        scan_dir = settings.DOCUMENTS_DIR



    if os.path.exists(scan_dir):

        for fname in os.listdir(scan_dir):

            ext = os.path.splitext(fname)[1].lower()

            if ext in extra_extensions and fname not in indexed_filenames:

                file_path = os.path.join(scan_dir, fname)

                if os.path.isfile(file_path):

                    docs.append(fname)



    # 统一格式为纯文件名字符串（前端兼容）

    normalized_docs = []

    for doc in docs:

        if isinstance(doc, dict):

            normalized_docs.append(doc.get('filename', doc.get('name', str(doc))))

        else:

            normalized_docs.append(str(doc))

    docs = normalized_docs



    total = len(docs)

    # 分页

    start = (page - 1) * page_size

    end = start + page_size

    paginated = docs[start:end]

    return {

        "documents": paginated,

        "count": total,

        "total": total,

        "page": page,

        "page_size": page_size,

        "total_pages": (total + page_size - 1) // page_size,

    }







@router.get("/documents/stats", summary="获取知识库统计信息")
async def get_document_stats(
    agent_id: str = Query(None, description="智能体ID，为空时查全局知识库"),
    username: str = Depends(require_auth),
):
    """获取知识库的文档数量和文本片段总数（按智能体隔离）
    
    统计来源：
    1. ChromaDB 向量索引中的分块数
    2. 关键词索引中的条目数
    取两者中较大的值作为总数
    """
    # [用户隔离] agent_id 加用户名前缀
    agent_id = _user_agent_id(agent_id, username)

    if not agent_id:
        return {"total_documents": 0, "total_chunks": 0, "indexing_mode": "none"}
    
    total_chunks = 0
    indexing_mode = "none"
    
    # 1. 从 ChromaDB 获取分块数
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is not None:
        try:
            collection = vector_store._collection
            chunk_count = collection.count()
            if chunk_count > 0:
                total_chunks = chunk_count
                indexing_mode = "vector"
        except Exception as e:
            logger.warning(f"获取ChromaDB分块数失败: {e}")
    
    # 2. 从关键词索引获取条目数
    try:
        keyword_index = _load_keyword_index(agent_id)
        keyword_count = len(keyword_index)
        if keyword_count > total_chunks:
            total_chunks = keyword_count
            indexing_mode = "keyword"
    except Exception as e:
        logger.warning(f"获取关键词索引条目数失败: {e}")
    
    # 3. 获取文档数量
    docs = list_indexed_documents(agent_id=agent_id)
    total_documents = len(docs)

    return {
        "total_documents": total_documents,
        "total_chunks": total_chunks,
        "indexing_mode": indexing_mode,
    }


# ===== 二级子目录管理 API（三列布局用）=====

def _user_agent_id(agent_id: str, username: str) -> str:
    """[用户隔离] 把 agent_id 转成用户隔离的 ID
    - __external__ 不转换（全质知识库所有用户共享）
    - 其他 agent_id → {username}_{agent_id}（每个用户的文件独立存储）"""
    if not agent_id or agent_id == "__external__":
        return agent_id
    return f"{username}_{agent_id}"

@router.get("/kb/categories", summary="列出所有一级分类")
async def list_categories_api(
    agent_id: str = Query(..., description="智能体ID"),
    username: str = Depends(require_auth),
):
    """列出指定智能体下所有一级分类名（从磁盘+ChromaDB+关键词索引合并）"""
    agent_id = _user_agent_id(agent_id, username)
    from app.rag.document import list_categories
    try:
        cats = await asyncio.to_thread(list_categories, agent_id)
        return {"success": True, "categories": cats}
    except Exception as e:
        logger.exception(f"列出分类失败: {e}")
        return {"success": True, "categories": ['手册', '程序文件', '三层次文件', '记录表格', '其他']}


@router.get("/kb/subcategories", summary="列出某一级分类下的二级子目录")
async def list_subcategories_api(
    agent_id: str = Query(..., description="智能体ID"),
    category: str = Query(..., description="一级分类名"),
    username: str = Depends(require_auth),
):
    """列出指定智能体的某一级分类下所有二级子目录名"""
    agent_id = _user_agent_id(agent_id, username)
    from app.rag.document import list_subcategories
    try:
        subcats = await asyncio.to_thread(list_subcategories, agent_id, category)
        return {"success": True, "subcategories": subcats}
    except Exception as e:
        logger.exception(f"列出子目录失败: {e}")
        return {"success": True, "subcategories": []}


# ===== 三级子目录管理 API（全质知识库用）=====

@router.get("/kb/subsubcategories", summary="列出某二级子目录下的三级子目录")
async def list_subsubcategories_api(
    agent_id: str = Query(..., description="智能体ID"),
    category: str = Query(..., description="一级分类名"),
    subcategory: str = Query(..., description="二级子目录名"),
    username: str = Depends(require_auth),
):
    """列出指定智能体的某二级子目录下所有三级子目录名"""
    agent_id = _user_agent_id(agent_id, username)
    from app.rag.document import list_subsubcategories
    try:
        subsubcats = await asyncio.to_thread(list_subsubcategories, agent_id, category, subcategory)
        return {"success": True, "subsubcategories": subsubcats}
    except Exception as e:
        logger.exception(f"列出三级子目录失败: {e}")
        return {"success": True, "subsubcategories": []}


@router.post("/kb/subsubcategories", summary="新建三级子目录")
async def create_subsubcategory_api(
    request: Request,
    username: str = Depends(require_auth),
):
    """新建三级子目录（在指定 agent_id + category + subcategory 下创建子文件夹）"""
    try:
        body = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体不是有效 JSON: {e}")
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    agent_id = (body.get("agent_id") or "").strip()
    agent_id = _user_agent_id(agent_id, username)
    category = (body.get("category") or "").strip()
    subcategory = (body.get("subcategory") or "").strip()
    subsubcategory = (body.get("subsubcategory") or "").strip()
    if not all([agent_id, category, subcategory, subsubcategory]):
        raise HTTPException(status_code=400, detail="agent_id/category/subcategory/subsubcategory 不能为空")
    import re as _re
    if _re.search(r'[\\/:*?"<>|]', subsubcategory):
        raise HTTPException(status_code=400, detail="子目录名含非法字符")
    subsub_dir = os.path.join(_get_agent_dir(agent_id), category, subcategory, subsubcategory)
    try:
        os.makedirs(subsub_dir, exist_ok=True)
    except Exception as e:
        logger.exception(f"[KB] 新建三级子目录失败: {subsub_dir}, err={e}")
        raise HTTPException(status_code=500, detail=f"创建目录失败: {e}")
    logger.info(f"[KB] 新建三级子目录: agent={agent_id}, cat={category}, sub={subcategory}, subsub={subsubcategory}, user={username}")
    return {"success": True, "subsubcategory": subsubcategory}


@router.put("/kb/subsubcategories", summary="重命名三级子目录")
async def rename_subsubcategory_api(
    request: Request,
    username: str = Depends(require_auth),
):
    """重命名三级子目录（同步更新磁盘/ChromaDB/关键词索引）"""
    from app.rag.document import rename_subsubcategory
    try:
        body = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体不是有效 JSON: {e}")
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    agent_id = (body.get("agent_id") or "").strip()
    agent_id = _user_agent_id(agent_id, username)
    category = (body.get("category") or "").strip()
    subcategory = (body.get("subcategory") or "").strip()
    old_subsub = (body.get("old_subsubcategory") or "").strip()
    new_subsub = (body.get("new_subsubcategory") or "").strip()
    if not all([agent_id, category, subcategory, old_subsub, new_subsub]):
        raise HTTPException(status_code=400, detail="参数不完整")
    result = await asyncio.to_thread(rename_subsubcategory, agent_id, category, subcategory, old_subsub, new_subsub)
    if result.get("status") != "success":
        raise HTTPException(status_code=400, detail=result.get("message", "重命名失败"))
    logger.info(f"[KB] 重命名三级子目录: {old_subsub} -> {new_subsub}, user={username}")
    return {"success": True, "message": result.get("message")}


@router.delete("/kb/subsubcategories", summary="删除三级子目录及其下所有文件")
async def delete_subsubcategory_api(
    request: Request,
    username: str = Depends(require_auth),
):
    """删除三级子目录（同时删除磁盘文件 + ChromaDB 文档 + 关键词索引）"""
    from app.rag.document import delete_subsubcategory
    try:
        body = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体不是有效 JSON: {e}")
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    agent_id = (body.get("agent_id") or "").strip()
    agent_id = _user_agent_id(agent_id, username)
    category = (body.get("category") or "").strip()
    subcategory = (body.get("subcategory") or "").strip()
    subsubcategory = (body.get("subsubcategory") or "").strip()
    if not all([agent_id, category, subcategory, subsubcategory]):
        raise HTTPException(status_code=400, detail="参数不完整")
    # 全质知识库（external）仅管理员可删除，普通用户只读+上传
    if agent_id == "__external__" and not is_admin(username):
        raise HTTPException(status_code=403, detail="权限不足，全质知识库仅管理员可删除")
    result = await asyncio.to_thread(delete_subsubcategory, agent_id, category, subcategory, subsubcategory)
    if result.get("status") != "success":
        raise HTTPException(status_code=400, detail=result.get("message", "删除失败"))
    logger.info(f"[KB] 删除三级子目录: {category}/{subcategory}/{subsubcategory}, user={username}")
    return {"success": True, "message": result.get("message")}


@router.post("/kb/subcategories", summary="新建二级子目录")
async def create_subcategory_api(
    request: Request,
    username: str = Depends(require_auth),
):
    """新建二级子目录（在指定 agent_id + category 下创建子文件夹）"""
    from app.rag.document import list_subcategories
    try:
        body = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体不是有效 JSON: {e}")
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    agent_id = (body.get("agent_id") or "").strip()
    agent_id = _user_agent_id(agent_id, username)
    category = (body.get("category") or "").strip()
    subcategory = (body.get("subcategory") or "").strip()
    if not agent_id or not category or not subcategory:
        raise HTTPException(status_code=400, detail="agent_id/category/subcategory 不能为空")
    # 校验名称不含非法字符
    import re as _re
    if _re.search(r'[\\/:*?"<>|]', subcategory):
        raise HTTPException(status_code=400, detail="子目录名含非法字符")
    # 检查是否已存在
    existing = await asyncio.to_thread(list_subcategories, agent_id, category)
    if subcategory in existing:
        raise HTTPException(status_code=400, detail=f"子目录「{subcategory}」已存在")
    # 创建磁盘文件夹（os.makedirs 会自动创建中间目录 agent_xxx/category）
    sub_dir = os.path.join(_get_agent_dir(agent_id), category, subcategory)
    try:
        os.makedirs(sub_dir, exist_ok=True)
    except Exception as e:
        logger.exception(f"[KB] 新建子目录失败: {sub_dir}, err={e}")
        raise HTTPException(status_code=500, detail=f"创建目录失败: {e}")
    logger.info(f"[KB] 新建子目录: agent={agent_id}, cat={category}, sub={subcategory}, user={username}, path={sub_dir}")

    # [需求实现] 程序文件的部门子目录同步到三层次文件和记录表格
    # 在程序文件里新建部门时，自动在三层次文件和记录表格里也新建同名部门
    SYNC_CATEGORIES = ["三层次文件", "记录表格"]
    synced = []
    if category == "程序文件" and agent_id != "__external__":
        for sync_cat in SYNC_CATEGORIES:
            try:
                sync_dir = os.path.join(_get_agent_dir(agent_id), sync_cat, subcategory)
                os.makedirs(sync_dir, exist_ok=True)
                synced.append(sync_cat)
            except Exception as sync_e:
                logger.warning(f"[KB] 同步新建子目录失败: {sync_cat}/{subcategory}, err={sync_e}")
        if synced:
            logger.info(f"[KB] 部门「{subcategory}」已同步到: {', '.join(synced)}")

    return {"success": True, "subcategory": subcategory, "synced_to": synced if synced else None}


@router.put("/kb/subcategories", summary="重命名二级子目录")
async def rename_subcategory_api(
    request: Request,
    username: str = Depends(require_auth),
):
    """重命名二级子目录（同步更新磁盘/ChromaDB/关键词索引）"""
    from app.rag.document import rename_subcategory
    try:
        body = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体不是有效 JSON: {e}")
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    agent_id = (body.get("agent_id") or "").strip()
    agent_id = _user_agent_id(agent_id, username)
    category = (body.get("category") or "").strip()
    old_sub = (body.get("old_subcategory") or "").strip()
    new_sub = (body.get("new_subcategory") or "").strip()
    if not all([agent_id, category, old_sub, new_sub]):
        raise HTTPException(status_code=400, detail="参数不完整")
    result = await asyncio.to_thread(rename_subcategory, agent_id, category, old_sub, new_sub)
    if result.get("status") != "success":
        raise HTTPException(status_code=400, detail=result.get("message", "重命名失败"))
    logger.info(f"[KB] 重命名子目录: {old_sub} -> {new_sub}, user={username}")

    # [需求实现] 程序文件的部门重命名同步到三层次文件和记录表格
    SYNC_CATEGORIES = ["三层次文件", "记录表格"]
    synced = []
    if category == "程序文件" and agent_id != "__external__":
        for sync_cat in SYNC_CATEGORIES:
            try:
                sync_result = await asyncio.to_thread(rename_subcategory, agent_id, sync_cat, old_sub, new_sub)
                if sync_result.get("status") == "success":
                    synced.append(sync_cat)
            except Exception as sync_e:
                # 三层次文件/记录表格里可能没有这个部门，忽略错误
                logger.debug(f"[KB] 同步重命名子目录跳过: {sync_cat}/{old_sub}, err={sync_e}")
        if synced:
            logger.info(f"[KB] 部门「{old_sub}→{new_sub}」已同步到: {', '.join(synced)}")

    return {"success": True, "message": result.get("message"), "synced_to": synced if synced else None}


@router.delete("/kb/subcategories", summary="删除二级子目录及其下所有文件")
async def delete_subcategory_api(
    request: Request,
    username: str = Depends(require_auth),
):
    """删除二级子目录（同时删除磁盘文件 + ChromaDB 文档 + 关键词索引）"""
    from app.rag.document import delete_subcategory
    try:
        body = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体不是有效 JSON: {e}")
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    agent_id = (body.get("agent_id") or "").strip()
    agent_id = _user_agent_id(agent_id, username)
    category = (body.get("category") or "").strip()
    subcategory = (body.get("subcategory") or "").strip()
    if not all([agent_id, category, subcategory]):
        raise HTTPException(status_code=400, detail="参数不完整")
    # 全质知识库（external）仅管理员可删除，普通用户只读+上传
    if agent_id == "__external__" and not is_admin(username):
        raise HTTPException(status_code=403, detail="权限不足，全质知识库仅管理员可删除")
    result = await asyncio.to_thread(delete_subcategory, agent_id, category, subcategory)
    if result.get("status") != "success":
        raise HTTPException(status_code=400, detail=result.get("message", "删除失败"))
    logger.info(f"[KB] 删除子目录: {category}/{subcategory}, user={username}")

    # [需求实现] 程序文件的部门删除同步到三层次文件和记录表格
    SYNC_CATEGORIES = ["三层次文件", "记录表格"]
    synced = []
    if category == "程序文件" and agent_id != "__external__":
        for sync_cat in SYNC_CATEGORIES:
            try:
                sync_result = await asyncio.to_thread(delete_subcategory, agent_id, sync_cat, subcategory)
                if sync_result.get("status") == "success":
                    synced.append(sync_cat)
            except Exception as sync_e:
                # 三层次文件/记录表格里可能没有这个部门，忽略错误
                logger.debug(f"[KB] 同步删除子目录跳过: {sync_cat}/{subcategory}, err={sync_e}")
        if synced:
            logger.info(f"[KB] 部门「{subcategory}」已从以下分类同步删除: {', '.join(synced)}")

    return {"success": True, "message": result.get("message"), "synced_to": synced if synced else None}


@router.put("/kb/categories", summary="重命名一级分类")
async def rename_category_api(
    request: Request,
    username: str = Depends(require_auth),
):
    """重命名一级分类：重命名磁盘文件夹 + 更新所有文档的 category metadata"""
    try:
        body = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体不是有效 JSON: {e}")
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    agent_id = (body.get("agent_id") or "").strip()
    agent_id = _user_agent_id(agent_id, username)
    old_cat = (body.get("old_category") or "").strip()
    new_cat = (body.get("new_category") or "").strip()
    if not all([agent_id, old_cat, new_cat]):
        raise HTTPException(status_code=400, detail="参数不完整")
    import re as _re
    if _re.search(r'[\\/:*?"<>|]', new_cat):
        raise HTTPException(status_code=400, detail="分类名含非法字符")

    # 1. 重命名磁盘文件夹
    old_dir = os.path.join(_get_agent_dir(agent_id), old_cat)
    new_dir = os.path.join(_get_agent_dir(agent_id), new_cat)
    if os.path.exists(new_dir):
        raise HTTPException(status_code=400, detail=f"分类「{new_cat}」已存在")
    if os.path.exists(old_dir):
        os.rename(old_dir, new_dir)

    # 2. 更新 ChromaDB metadata
    from app.rag.document import get_vector_store, _load_keyword_index, _save_keyword_index, _bm25_cache_invalidation
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is not None:
        try:
            collection = vector_store._collection
            all_docs = collection.get(include=["metadatas"])
            for i, meta in enumerate(all_docs["metadatas"]):
                if meta and meta.get("category") == old_cat:
                    doc_id = all_docs["ids"][i]
                    new_meta = dict(meta)
                    new_meta["category"] = new_cat
                    collection.update(ids=[doc_id], metadatas=[new_meta])
        except Exception as e:
            logger.warning(f"更新 ChromaDB category 失败: {e}")

    # 3. 更新关键词索引
    keyword_docs = _load_keyword_index(agent_id)
    updated = False
    for entry in keyword_docs:
        if entry.get("category") == old_cat:
            entry["category"] = new_cat
            updated = True
    if updated:
        _save_keyword_index(keyword_docs, agent_id)
        _bm25_cache_invalidation(agent_id)

    logger.info(f"[KB] 重命名一级分类: {old_cat} -> {new_cat}, user={username}")
    return {"success": True, "message": f"已重命名「{old_cat}」→「{new_cat}」"}


@router.delete("/kb/categories", summary="删除一级分类及其下所有子目录和文件")
async def delete_category_api(
    request: Request,
    username: str = Depends(require_auth),
):
    """删除一级分类：删除磁盘文件夹 + ChromaDB 文档 + 关键词索引"""
    try:
        body = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体不是有效 JSON: {e}")
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    agent_id = (body.get("agent_id") or "").strip()
    agent_id = _user_agent_id(agent_id, username)
    category = (body.get("category") or "").strip()
    if not all([agent_id, category]):
        raise HTTPException(status_code=400, detail="参数不完整")
    # 全质知识库（external）仅管理员可删除，普通用户只读+上传
    if agent_id == "__external__" and not is_admin(username):
        raise HTTPException(status_code=403, detail="权限不足，全质知识库仅管理员可删除")

    # 1. 删除磁盘文件夹
    cat_dir = os.path.join(_get_agent_dir(agent_id), category)
    if os.path.exists(cat_dir):
        import shutil
        shutil.rmtree(cat_dir, ignore_errors=True)

    # 2. 从 ChromaDB 删除
    from app.rag.document import get_vector_store, _load_keyword_index, _save_keyword_index, _bm25_cache_invalidation
    vector_store = get_vector_store(agent_id=agent_id)
    if vector_store is not None:
        try:
            collection = vector_store._collection
            all_docs = collection.get(include=["metadatas"])
            ids_to_delete = []
            for i, meta in enumerate(all_docs["metadatas"]):
                if meta and meta.get("category") == category:
                    ids_to_delete.append(all_docs["ids"][i])
            if ids_to_delete:
                collection.delete(ids=ids_to_delete)
        except Exception as e:
            logger.warning(f"从 ChromaDB 删除分类文档失败: {e}")

    # 3. 从关键词索引删除
    keyword_docs = _load_keyword_index(agent_id)
    new_keyword_docs = [e for e in keyword_docs if e.get("category") != category]
    if len(new_keyword_docs) != len(keyword_docs):
        _save_keyword_index(new_keyword_docs, agent_id)
        _bm25_cache_invalidation(agent_id)

    logger.info(f"[KB] 删除一级分类: {category}, user={username}")
    return {"success": True, "message": f"已删除分类「{category}」及其下所有文件"}


@router.post("/kb/categories", summary="新建一级分类")
async def create_category_api(
    request: Request,
    username: str = Depends(require_auth),
):
    """新建一级分类（在 data/documents/agent_{agent_id}/ 下创建文件夹）"""
    try:
        body = await request.json()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"请求体不是有效 JSON: {e}")
    if not isinstance(body, dict):
        raise HTTPException(status_code=400, detail="请求体必须是 JSON 对象")
    agent_id = (body.get("agent_id") or "").strip()
    agent_id = _user_agent_id(agent_id, username)
    category = (body.get("category") or "").strip()
    if not agent_id or not category:
        raise HTTPException(status_code=400, detail="agent_id 和 category 不能为空")
    import re as _re
    if _re.search(r'[\\/:*?"<>|]', category):
        raise HTTPException(status_code=400, detail="分类名含非法字符")
    cat_dir = os.path.join(_get_agent_dir(agent_id), category)
    if os.path.exists(cat_dir):
        raise HTTPException(status_code=400, detail=f"分类「{category}」已存在")
    try:
        os.makedirs(cat_dir, exist_ok=True)
    except Exception as e:
        logger.exception(f"[KB] 新建一级分类失败: {cat_dir}, err={e}")
        raise HTTPException(status_code=500, detail=f"创建目录失败: {e}")
    logger.info(f"[KB] 新建一级分类: {category}, user={username}, path={cat_dir}")
    return {"success": True, "category": category}


@router.put("/documents/{filename}", summary="修改知识库文档内容")

async def modify_document_api(filename: str, req: ModifyDocumentRequest, username: str = Depends(require_auth)):

    """

    修改知识库中指定文档的内容

    支持两种模式：

    - 替换模式（append=false）：用新内容完全替换原文档内容

    - 追加模式（append=true）：在原文档内容末尾追加新内容

    修改后会自动重新索引到向量数据库

    """

    # 检查文档是否存在（优先查找agent子目录）

    file_path = None

    if req.agent_id:

        agent_path = os.path.join(settings.DOCUMENTS_DIR, f"agent_{_user_agent_id(req.agent_id, username)}", filename)

        if os.path.exists(agent_path):

            file_path = agent_path

    if not file_path:

        global_path = os.path.join(settings.DOCUMENTS_DIR, filename)

        if os.path.exists(global_path):

            file_path = global_path

    if not file_path:

        raise HTTPException(status_code=404, detail=f"文档 {filename} 不存在")



    # 追加模式：先读取原内容，拼接新内容

    final_content = req.content

    if req.append:

        try:

            from app.rag.document import load_document

            docs = await asyncio.to_thread(load_document, file_path)

            original_text = "\n".join([doc.page_content for doc in docs])

            final_content = original_text + "\n" + req.content

        except Exception as e:

            logger.warning(f"读取原文档内容失败，改为替换模式: {e}")



    logger.info(f"知识库修改文档: {filename}, 追加模式={req.append}, 内容长度={len(final_content)}, agent_id={req.agent_id}")



    result = update_document(filename, final_content, agent_id=_user_agent_id(req.agent_id, username), async_reindex=True)  # 异步重索引，加速响应

    if result["status"] == "not_found":

        raise HTTPException(status_code=404, detail=result["message"])

    if result["status"] == "error":

        raise HTTPException(status_code=500, detail=result["message"])



    response_data = {"status": "success", "detail": result}



    # 如果用户要求返回docx文件下载链接

    if req.return_docx:

        try:

            docx_filename = filename.rsplit('.', 1)[0] + '.docx'

            docx_result = export_document_as_docx(final_content, docx_filename, session_id=req.session_id or "")

            if docx_result["status"] == "success":

                actual_docx_filename = docx_result.get('filename', docx_filename)

                response_data["download_url"] = f"/api/v1/documents/export-download/{actual_docx_filename}?sid={req.session_id or ''}"

                response_data["docx_filename"] = actual_docx_filename

        except Exception as e:

            logger.warning(f"生成docx下载文件失败: {e}")



    return response_data





@router.get("/documents/{filename}/download", summary="下载知识库文档")

async def download_document(filename: str, agent_id: str = Query(None, description="智能体ID，为空时查全局知识库"), auth_username: str = Depends(require_auth)):

    """

    下载知识库中的文档文件

    支持 .docx / .txt / .pdf 格式
    [Bug 3 修复] 加 JWT 认证 + agent_id 白名单 + 路径穿越校验
    """

    # [Bug 3 修复] agent_id 白名单校验，防止路径穿越
    if agent_id:
        if not re.match(r'^[A-Za-z0-9_-]+$', agent_id or ''):
            raise HTTPException(status_code=400, detail="非法的 agent_id")

    # [用户隔离] agent_id 加用户名前缀
    agent_id = _user_agent_id(agent_id, auth_username)
    
    # [Bug 3 修复] filename 白名单校验，防止路径穿越
    if not re.match(r'^[A-Za-z0-9_\u4e00-\u9fa5.\-() ]+$', filename or '') or '..' in filename or '/' in filename or '\\' in filename:
        raise HTTPException(status_code=400, detail="非法的文件名")

    # 先查找agent子目录，再查全局目录

    if agent_id:

        file_path = os.path.join(settings.DOCUMENTS_DIR, f"agent_{agent_id}", filename)

    else:

        file_path = os.path.join(settings.DOCUMENTS_DIR, filename)

    # [Bug 3 修复] 校验最终路径仍在 DOCUMENTS_DIR 下（防止符号链接等绕过）
    real_path = os.path.realpath(file_path)
    real_docs_dir = os.path.realpath(settings.DOCUMENTS_DIR)
    if not real_path.startswith(real_docs_dir + os.sep) and real_path != real_docs_dir:
        raise HTTPException(status_code=400, detail="非法的文件路径")

    if not os.path.exists(file_path):

        # 回退：尝试全局目录

        fallback_path = os.path.join(settings.DOCUMENTS_DIR, filename)

        if os.path.exists(fallback_path):

            file_path = fallback_path

        else:

            raise HTTPException(status_code=404, detail=f"文档 {filename} 不存在")



    ext = os.path.splitext(filename)[1].lower()

    mime_map = {

        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",

        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",

        ".pdf": "application/pdf",

        ".txt": "text/plain; charset=utf-8",

    }

    media_type = mime_map.get(ext, "application/octet-stream")



    with open(file_path, "rb") as f:

        content = f.read()



    # RFC 5987: 中文文件名需要URL编码

    from urllib.parse import quote

    encoded_filename = quote(filename)



    return Response(

        content=content,

        media_type=media_type,

        headers={

            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"

        }

    )





@router.post("/documents/export", summary="导出/生成文档为docx")

async def export_document_api(req: ExportDocumentRequest):

    """

    将文本内容生成为docx文档并提供下载

    支持从知识库内容整合生成综合文档或简略文档

    """

    try:

        filename = req.filename or f"export_{int(time.time())}.docx"

        if not filename.endswith('.docx'):

            filename += '.docx'



        result = export_document_as_docx(req.content, filename, title=req.title, session_id=req.session_id or "")

        if result["status"] == "success":

            actual_filename = result.get('filename', filename)

            return {

                "status": "success",

                "filename": actual_filename,

                "download_url": f"/api/v1/documents/export-download/{actual_filename}?sid={req.session_id or ''}",

                "message": result["message"],

            }

        else:

            raise HTTPException(status_code=500, detail=result.get("message", "导出失败"))

    except Exception as e:

        raise HTTPException(status_code=500, detail=f"文档导出失败: {str(e)}")





class ExportXlsxRequest(BaseModel):

    """导出/生成XLSX文档请求"""

    content: str  # 文档内容（Markdown格式，支持表格）

    filename: str = ""  # 输出文件名（含扩展名），为空则自动生成

    title: str = ""  # 文档标题/工作表名称，为空则使用filename

    session_id: str = ""  # [Bug 修复] 会话ID，用于导出文件按会话隔离目录





@router.post("/documents/export-xlsx", summary="导出/生成文档为xlsx")

async def export_xlsx_api(req: ExportXlsxRequest):

    """

    将文本内容生成为xlsx（Excel）文档并提供下载

    支持Markdown表格自动转为Excel原生表格

    """

    try:

        from app.rag.document import export_document_as_xlsx

        

        filename = req.filename or f"export_{int(time.time())}.xlsx"

        if not filename.endswith('.xlsx'):

            filename = filename.rsplit('.', 1)[0] + '.xlsx'



        result = export_document_as_xlsx(req.content, filename, title=req.title, session_id=req.session_id or "")

        if result["status"] == "success":

            actual_filename = result.get('filename', filename)

            return {

                "status": "success",

                "filename": actual_filename,

                "download_url": f"/api/v1/documents/export-download/{actual_filename}?sid={req.session_id or ''}",

                "message": result["message"],

            }

        else:

            raise HTTPException(status_code=500, detail=result.get("message", "导出失败"))

    except Exception as e:

        raise HTTPException(status_code=500, detail=f"文档导出失败: {str(e)}")





@router.get("/documents/export-download/{filename}", summary="下载AI导出的文档")

async def download_export_document(filename: str, sid: str = None):

    """

    下载AI生成的导出文档（docx/txt）

    文件保存在 data/export/{session_id}/ 目录中

    支持会话子目录查找 + 兼容旧版平铺目录

    [Bug 修复] 新增 sid 查询参数，优先在指定会话目录中查找，避免同名文件下载到其他会话的版本

    """

    from urllib.parse import unquote

    import unicodedata



    # URL解码文件名（处理中文文件名）

    # FastAPI可能已经自动解码一次，再unquote确保双重编码也能处理

    decoded_filename = unquote(unquote(filename))

    # 安全检查：防止路径穿越

    safe_filename = decoded_filename.replace('/', '_').replace('\\', '_').replace('..', '_')



    # sid 也做安全检查（防止路径穿越）
    safe_sid = (sid or '').replace('/', '_').replace('\\', '_').replace('..', '_') if sid else ''

    export_root = _get_export_dir()  # export 根目录

    file_path = None



    # 0. [Bug 修复] 优先在 sid 指定的会话目录中查找，避免同名文件下载到其他会话的版本
    if safe_sid:
        candidate_dir = os.path.join(export_root, safe_sid)
        if os.path.isdir(candidate_dir):
            candidate = os.path.join(candidate_dir, safe_filename)
            if os.path.exists(candidate):
                file_path = candidate
                logger.info(f"[导出下载] 在指定会话目录 {safe_sid}/ 中找到文件: {safe_filename}")



    # 1. 如果 sid 没找到，遍历所有会话子目录查找（兼容旧版无 sid 的链接）

    if file_path is None and os.path.exists(export_root):

        for item in os.listdir(export_root):

            sub_dir = os.path.join(export_root, item)

            if os.path.isdir(sub_dir):

                candidate = os.path.join(sub_dir, safe_filename)

                if os.path.exists(candidate):

                    file_path = candidate

                    logger.info(f"[导出下载] 在会话目录 {item}/ 中找到文件: {safe_filename}")

                    break



    # 2. 兼容旧版：直接在 export 根目录查找

    if file_path is None:

        file_path = os.path.join(export_root, safe_filename)



    # 3. 精确匹配

    if os.path.exists(file_path):

        logger.info(f"[导出下载] 文件匹配成功: {safe_filename}")

    else:

        # 4. 模糊匹配：尝试Unicode标准化 + 不带扩展名匹配

        found = False



        # 方法1：NFC/NFD Unicode标准化

        norm_filename = unicodedata.normalize('NFC', safe_filename)

        # 先搜子目录

        if os.path.exists(export_root):

            for item in os.listdir(export_root):

                search_dir = os.path.join(export_root, item) if os.path.isdir(os.path.join(export_root, item)) else export_root

                norm_path = os.path.join(search_dir, norm_filename)

                if os.path.exists(norm_path):

                    file_path = norm_path

                    safe_filename = norm_filename

                    found = True

                    logger.info(f"[导出下载] 通过Unicode标准化匹配成功: {safe_filename}")

                    break



        # 方法2：遍历所有目录做模糊匹配（忽略Unicode差异）

        if not found and os.path.exists(export_root):

            base_name = os.path.splitext(safe_filename)[0]

            # 遍历根目录和所有子目录

            search_dirs = [export_root]

            for item in os.listdir(export_root):

                sub = os.path.join(export_root, item)

                if os.path.isdir(sub):

                    search_dirs.append(sub)



            for search_dir in search_dirs:

                if not os.path.exists(search_dir):

                    continue

                for existing_file in os.listdir(search_dir):

                    existing_path = os.path.join(search_dir, existing_file)

                    if not os.path.isfile(existing_path):

                        continue

                    existing_base = os.path.splitext(existing_file)[0]

                    # 比较Unicode标准化后的文件名

                    if (unicodedata.normalize('NFC', existing_base) == unicodedata.normalize('NFC', base_name)

                        and os.path.splitext(safe_filename)[1].lower() == os.path.splitext(existing_file)[1].lower()):

                        file_path = existing_path

                        safe_filename = existing_file

                        found = True

                        logger.info(f"[导出下载] 通过模糊匹配找到文件: {existing_file} (请求: {decoded_filename})")

                        break

                if found:

                    break



        if not found:

            # 记录目录中现有文件，帮助调试

            existing_files = []

            if os.path.exists(export_root):

                for item in os.listdir(export_root):

                    sub = os.path.join(export_root, item)

                    if os.path.isdir(sub):

                        existing_files.extend([f"{item}/{f}" for f in os.listdir(sub) if os.path.isfile(os.path.join(sub, f))])

                    elif os.path.isfile(sub):

                        existing_files.append(item)

            logger.warning(f"[导出下载] 文件未找到! 请求文件名: {safe_filename}, 目录中现有文件: {existing_files[:10]}")

            raise HTTPException(status_code=404, detail=f"导出文档 {safe_filename} 不存在")



    ext = os.path.splitext(safe_filename)[1].lower()

    mime_map = {

        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",

        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",

        ".pdf": "application/pdf",

        ".txt": "text/plain; charset=utf-8",

    }

    media_type = mime_map.get(ext, "application/octet-stream")



    # 使用 RFC 5987 编码处理中文文件名
    from urllib.parse import quote
    encoded_filename = quote(safe_filename)

    return FileResponse(
        path=file_path,
        media_type=media_type,
        filename=safe_filename,
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )






@router.delete("/documents/{filename}", summary="从知识库删除文档")

async def delete_document_api(filename: str, agent_id: str = Query(None, description="智能体ID，为空时删全局知识库文档"), username: str = Depends(require_auth)):

    """

    从知识库中删除指定文档（登录用户可操作自己的知识库）

    同时删除 ChromaDB 中的向量分块和原始文件



    注意：普通聊天模式（agent_id=None）没有知识库，不支持删除

    """

    # [用户隔离] agent_id 加用户名前缀
    agent_id = _user_agent_id(agent_id, username)

    # 普通聊天模式没有知识库

    if not agent_id:

        raise HTTPException(status_code=400, detail="普通聊天模式没有知识库，不支持删除文档。请先选择一个智能体。")

    result = delete_document(filename, agent_id=agent_id)

    if result["status"] == "not_found":

        raise HTTPException(status_code=404, detail=result["message"])

    if result["status"] == "error":

        raise HTTPException(status_code=500, detail=result["message"])

    return {"status": "success", "detail": result}





# ===== 会话历史接口 =====



@router.get("/history/{session_id}", summary="获取对话历史")

async def get_history(session_id: str, auth_username: str = Depends(require_auth)):

    """获取指定会话的对话历史
    [Bug 2 修复] 加 JWT 认证"""

    # 校验会话归属：session_id 应以用户名开头
    if not session_id.startswith(auth_username + "_"):
        raise HTTPException(status_code=403, detail="无权访问该会话")

    messages = get_history_messages(session_id)

    return {"session_id": session_id, "messages": messages, "count": len(messages)}



@router.post("/history/{session_id}", summary="添加AI消息到对话历史（不走AI）")
async def add_ai_message_to_history(session_id: str, request: Request, auth_username: str = Depends(require_auth)):
    """直接添加一条 AI 消息到会话历史（不走AI，用于付费提示等固定回复）
    前端检测消息中的标记重新渲染对应样式"""
    if not session_id.startswith(auth_username + "_"):
        raise HTTPException(status_code=403, detail="无权操作该会话")
    body = await request.json()
    message = body.get("message", "")
    if not message:
        raise HTTPException(status_code=400, detail="message 不能为空")
    from langchain_core.messages import AIMessage
    from app.memory.manager import get_session_history, flush_session
    history = get_session_history(session_id)
    history.add_message(AIMessage(content=message))
    parts = session_id.split("_", 1)
    if len(parts) == 2:
        update_chat_time(parts[0], session_id)
    flush_session(session_id)
    return {"success": True}





@router.delete("/history/{session_id}", summary="清除对话历史")

async def delete_history(session_id: str, auth_username: str = Depends(require_auth)):

    """清除指定会话的对话历史，同时清理临时文件和导出文件
    [Bug 2 修复] 加 JWT 认证 + 会话归属校验"""

    # 校验会话归属
    if not session_id.startswith(auth_username + "_"):
        raise HTTPException(status_code=403, detail="无权操作该会话")

    clear_session_history(session_id)

    # 清理普通模式下的临时上传文件（data/temp/{session_id}/）

    temp_dir = os.path.join(settings.DATA_DIR, "temp", session_id)

    if os.path.exists(temp_dir):

        try:

            shutil.rmtree(temp_dir)

            logger.info(f"清空历史时清理临时文件: {temp_dir}")

        except Exception as e:

            logger.warning(f"清理临时文件失败: {e}")

    # 清理该会话的导出文件（data/export/{session_id}/）

    try:

        deleted_count = cleanup_export_files(session_id=session_id)

        if deleted_count > 0:

            logger.info(f"清空历史时清理了 {deleted_count} 个导出文件")

    except Exception as e:

        logger.warning(f"清理导出文件失败: {e}")

    return {"status": "success", "message": f"会话 {session_id} 的历史已清除"}





# ===== 会话管理接口 =====



@router.get("/chats", summary="获取用户会话列表")

async def get_chats(

    auth_username: str = Depends(require_auth),

    username: str = Query(None, description="（已废弃，统一用JWT认证）历史兼容参数"),

    mode: str = Query(None, description="模式过滤: agent/chat"),

    page: int = Query(1, ge=1, description="页码"),          # [#23] 分页

    page_size: int = Query(20, ge=1, le=100, description="每页数量"),

):

    """获取用户的会话列表（支持分页，支持按模式过滤）
    [Bug 2 修复] 从 JWT 拿 username，不再信任 query 参数"""

    username = auth_username  # 以 JWT 认证的用户名为准

    chats = list_chats(username, mode=mode, skip_auto_title=True)  # GET请求跳过自动标题更新，避免写副作用

    total = len(chats)

    start = (page - 1) * page_size

    end = start + page_size

    paginated = chats[start:end]

    return {

        "success": True,

        "chats": paginated,

        "total": total,

        "page": page,

        "page_size": page_size,

    }





@router.post("/chats", summary="创建新会话")

async def create_chat_api(

    auth_username: str = Depends(require_auth),

    title: str = "新对话",

    mode: str = "agent",

    agent_id: str = Query(None, description="智能体ID，会话归属到指定智能体"),

    username: str = Query(None, description="（已废弃）历史兼容参数"),

):

    """为用户创建一个新的会话（支持指定模式和智能体归属）
    [Bug 2 修复] 从 JWT 拿 username"""

    username = auth_username

    chat_info = create_chat(username, title, mode=mode, agent_id=agent_id)

    record_session()

    return {"success": True, "chat": chat_info}





@router.delete("/chats/{chat_id}", summary="删除会话")

async def delete_chat_api(chat_id: str, auth_username: str = Depends(require_auth)):

    """删除用户的某个会话，同时清理普通模式下的临时文件和导出文件
    [Bug 2 修复] 从 JWT 拿 username"""

    username = auth_username

    delete_chat(username, chat_id)

    # 清理普通模式下的临时文件（存放在 data/temp/{session_id}/ 目录）

    temp_dir = os.path.join(settings.DATA_DIR, "temp", chat_id)

    if os.path.exists(temp_dir):

        try:

            shutil.rmtree(temp_dir)

            logger.info(f"已清理临时文件: {temp_dir}")

        except Exception as e:

            logger.warning(f"清理临时文件失败: {e}")

    # 清理该会话的导出文件（data/export/{chat_id}/，只删当前会话的）

    try:

        deleted_count = cleanup_export_files(session_id=chat_id)

        if deleted_count > 0:

            logger.info(f"已清理 {deleted_count} 个导出文件")

    except Exception as e:

        logger.warning(f"清理导出文件失败: {e}")

    return {"success": True, "message": "会话已删除"}





@router.put("/chats/{chat_id}/rename", summary="重命名会话")

async def rename_chat_api(chat_id: str, req: RenameRequest, auth_username: str = Depends(require_auth)):

    """重命名用户的某个会话
    [Bug 2 修复] 从 JWT 拿 username，校验会话归属"""

    username = auth_username  # 以 JWT 为准，忽略 req.body 里的 username

    # 校验会话归属
    if not chat_id.startswith(username + "_"):
        raise HTTPException(status_code=403, detail="无权操作该会话")

    rename_chat(username, req.chat_id, req.new_title)

    return {"success": True, "message": "会话已重命名"}





# ===== 模型管理接口 =====



@router.get("/models", summary="获取可用模型列表")

async def get_models():

    """获取所有可用的 LLM 模型列表"""

    current = get_current_model()

    return {"models": AVAILABLE_MODELS, "current": current}





@router.post("/models/set", summary="切换模型")

async def set_model(req: ModelSetRequest):

    """切换当前使用的 LLM 模型"""

    success = set_current_model(req.model_id)

    if success:

        return {"success": True, "message": f"已切换到模型: {req.model_id}"}

    return {"success": False, "message": f"不支持的模型: {req.model_id}"}





# ===== 技能列表接口 =====



@router.get("/skills", summary="获取可用技能列表")

async def get_skills():

    """返回所有可用技能的信息（供前端 Skills 下拉菜单使用）"""

    import os

    skills_list = []

    skills_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "skills")

    if os.path.isdir(skills_dir):

        for name in sorted(os.listdir(skills_dir)):

            skill_dir = os.path.join(skills_dir, name)

            skill_md = os.path.join(skill_dir, "SKILL.md")

            if os.path.isdir(skill_dir) and os.path.isfile(skill_md):

                # 读取 SKILL.md 的前两行提取 name 和 description

                try:

                    with open(skill_md, "r", encoding="utf-8") as f:

                        front = f.read(512)

                    skill_name = name

                    skill_desc = ""

                    for line in front.split("\n"):

                        if line.startswith("name:"):

                            skill_name = line.split(":", 1)[1].strip()

                        elif line.startswith("description:"):

                            skill_desc = line.split(":", 1)[1].strip().strip("\"")

                    skills_list.append({

                        "id": name,

                        "name": skill_name,

                        "description": skill_desc,

                    })

                except Exception:

                    skills_list.append({"id": name, "name": name, "description": ""})

    return {"success": True, "skills": skills_list}



# ===== 使用统计接口 =====



@router.get("/stats", summary="获取使用统计")

async def get_usage_stats(username: str = Depends(get_current_user)):

    """获取系统使用统计数据"""

    stats = get_stats()

    # [#20] 附加 API 性能指标

    stats["api_performance"] = {

        "total_requests": _request_stats["total_requests"],

        "total_errors": _request_stats["total_errors"],

        "avg_response_time_ms": round(_request_stats["avg_response_time"] * 1000, 2),

        "error_rate": round(_request_stats["total_errors"] / max(_request_stats["total_requests"], 1) * 100, 2),

    }

    return {"success": True, "stats": stats}





# ===== [#22] 配置中心 API =====



@router.get("/config", summary="获取运行时配置")

async def get_config(username: str = Depends(require_auth)):

    """获取当前运行时配置（隐藏敏感信息）"""

    return {

        "success": True,

        "config": {

            "LLM_MODEL": settings.LLM_MODEL,

            "LLM_BASE_URL": settings.LLM_BASE_URL,

            "EMBEDDING_MODEL": settings.EMBEDDING_MODEL,

            "APP_HOST": settings.APP_HOST,

            "APP_PORT": settings.APP_PORT,

            "GITHUB_TOKEN_CONFIGURED": bool(os.getenv("GITHUB_TOKEN", "")),

            "SMTP_CONFIGURED": bool(os.getenv("SMTP_HOST", "")),

            "DATABASE_CONFIGURED": bool(os.getenv("DATABASE_URL", "")),

        }

    }





@router.post("/config", summary="更新运行时配置（热更新）")

async def update_config(req: ConfigUpdateRequest, username: str = Depends(require_auth)):

    """

    [#22] 运行时热更新配置，无需重启服务

    支持更新的配置项：LLM_MODEL, APP_PORT 等

    """

    allowed_keys = {"LLM_MODEL", "APP_PORT", "EMBEDDING_MODEL"}

    

    if req.key not in allowed_keys:

        raise HTTPException(status_code=400, detail=f"不允许更新的配置项: {req.key}。支持: {allowed_keys}")

    

    old_value = getattr(settings, req.key, None)

    if old_value is None:

        raise HTTPException(status_code=400, detail=f"未知的配置项: {req.key}")

    

    # 类型转换

    try:

        if req.key == "APP_PORT":

            new_value = int(req.value)

        else:

            new_value = req.value

    except ValueError:

        raise HTTPException(status_code=400, detail=f"配置值类型错误: {req.key} 期望 {type(old_value).__name__}")

    

    # 应用更新

    setattr(settings, req.key, new_value)

    

    # 如果更新了模型，重置 Agent

    if req.key == "LLM_MODEL":

        reset_agent()

        logger.info(f"配置热更新: {req.key} = {new_value}, Agent 已重置")

    elif req.key == "EMBEDDING_MODEL":

        from app.rag.document import reset_vector_store

        reset_vector_store()

        logger.info(f"配置热更新: {req.key} = {new_value}, 向量数据库已重置")

    

    logger.info(f"配置热更新: {req.key} 由 {old_value} 变更为 {new_value}, 操作者: {username}")

    

    return {

        "success": True,

        "message": f"配置 {req.key} 已更新",

        "old_value": str(old_value),

        "new_value": str(new_value),

    }





# ===== 导出对话接口 =====



@router.get("/export/{session_id}", summary="导出对话")

async def export_chat(session_id: str, format: str = "md", agent_name: str = "", auth_username: str = Depends(require_auth)):

    """

    导出对话为 Word(docx)、PDF 或 Markdown 格式
    [Bug 2 修复] 加 JWT 认证 + 会话归属校验

    format: docx | pdf | md
    agent_name: 当前智能体名称（用于文件名和标题）

    说明：
    - docx / pdf 都会先解析消息内容中的 Markdown，再渲染为对应格式的
      原生元素（Word/PDF 表格、标题、列表、代码块等），避免出现 |---|---|
      这样的纯文本残留。
    """

    # 校验会话归属
    if not session_id.startswith(auth_username + "_"):
        raise HTTPException(status_code=403, detail="无权导出该会话")

    # [BUG FIX] 使用 get_history_messages_from_file 强制从文件读取最新数据
    # 避免多worker进程下内存缓存不一致导致导出内容错误
    messages = get_history_messages_from_file(session_id)

    if not messages:

        raise HTTPException(status_code=404, detail="没有可导出的对话内容")



    # 安全文件名
    from urllib.parse import quote
    safe_name = agent_name.replace('/', '_').replace('\\', '_').replace('..', '_') if agent_name else f'chat_{session_id[:12]}'

    if format == "docx":

        # Word (docx) 导出 —— 解析 Markdown 后渲染为原生 Word 元素
        try:
            from app.utils.chat_export import generate_chat_docx_bytes
            docx_bytes = generate_chat_docx_bytes(messages, session_id, agent_name=agent_name)

            filename = f"{safe_name}_对话记录.docx"
            encoded_filename = quote(filename)

            return Response(

                content=docx_bytes,

                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",

                headers={

                    "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"

                }

            )

        except Exception as e:

            logger.error(f"[导出 Word] 失败: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"Word 生成失败: {str(e)}")



    elif format == "pdf":

        # PDF 导出 —— 同样解析 Markdown，表格用 reportlab Table 渲染
        try:

            from app.utils.chat_export import generate_chat_pdf_bytes
            pdf_bytes = generate_chat_pdf_bytes(messages, session_id, agent_name=agent_name)

            filename = f"{safe_name}_对话记录.pdf"
            encoded_filename = quote(filename)

            return Response(

                content=pdf_bytes,

                media_type="application/pdf",

                headers={

                    "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"

                }

            )

        except Exception as e:

            logger.error(f"[导出 PDF] 失败: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"PDF 生成失败: {str(e)}")

    else:

        # Markdown 导出
        display_title = f"{agent_name} 对话记录" if agent_name else "东风科技研发智能体 对话记录"
        content = f"# {display_title}\n\n"

        for msg in messages:

            role = "用户" if msg["role"] == "user" else "助手"

            content += f"**{role}：**\n\n{msg['content']}\n\n---\n\n"



        md_filename = f"{safe_name}_对话记录.md"
        encoded_md_filename = quote(md_filename)

        return Response(

            content=content.encode("utf-8"),

            media_type="text/markdown; charset=utf-8",

            headers={

                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_md_filename}"

            }

        )





# ===== [#24] 健康检查增强 =====



@router.get("/health/detailed", summary="详细健康检查")

async def health_detailed():

    """

    [#24] 详细健康检查：检查所有依赖组件状态

    - ChromaDB 可用性

    - LLM API 可达性

    - 磁盘空间

    - 内存使用

    """

    import platform

    

    checks = {}

    overall = "healthy"

    

    # 1. ChromaDB / 索引模式 检查

    indexing_mode = get_indexing_mode()

    if indexing_mode == "vector":

        try:

            from app.rag.document import get_vector_store

            vs = get_vector_store()

            if vs is not None:

                collection = vs._collection

                count = collection.count()

                checks["chromadb"] = {"status": "ok", "document_count": count, "indexing_mode": "vector"}

            else:

                checks["chromadb"] = {"status": "degraded", "indexing_mode": "keyword", "message": "Embedding 不可用，已自动降级为关键词搜索"}

        except Exception as e:

            checks["chromadb"] = {"status": "degraded", "indexing_mode": "keyword", "message": str(e)[:200]}

    elif indexing_mode == "keyword":

        checks["chromadb"] = {"status": "degraded", "indexing_mode": "keyword", "message": "Embedding API 不可用，已自动降级为关键词搜索模式"}

    else:

        checks["chromadb"] = {"status": "ok", "indexing_mode": "unknown", "message": "尚未检测 Embedding 可用性"}

    

    # 2. LLM API 检查

    try:

        import httpx

        api_url = settings.LLM_BASE_URL.rstrip("/") + "/models"

        resp = httpx.get(api_url, timeout=5)

        if resp.status_code == 200:

            checks["llm_api"] = {"status": "ok", "model": settings.LLM_MODEL}

        else:

            checks["llm_api"] = {"status": "error", "code": resp.status_code}

            overall = "degraded"

    except Exception as e:

        checks["llm_api"] = {"status": "unreachable", "message": str(e)[:100]}

        overall = "degraded"

    

    # 3. 磁盘空间检查

    try:

        disk_usage = shutil.disk_usage(settings.DATA_DIR)

        free_gb = disk_usage.free / (1024 ** 3)

        total_gb = disk_usage.total / (1024 ** 3)

        usage_pct = (disk_usage.used / disk_usage.total) * 100

        checks["disk"] = {

            "status": "ok" if usage_pct < 90 else "warning",

            "free_gb": round(free_gb, 2),

            "total_gb": round(total_gb, 2),

            "usage_percent": round(usage_pct, 1),

        }

        if usage_pct >= 90:

            overall = "degraded"

    except Exception as e:

        checks["disk"] = {"status": "error", "message": str(e)[:100]}

    

    # 4. 内存检查

    try:

        import psutil

        mem = psutil.virtual_memory()

        checks["memory"] = {

            "status": "ok" if mem.percent < 90 else "warning",

            "total_gb": round(mem.total / (1024 ** 3), 2),

            "used_percent": mem.percent,

        }

    except ImportError:

        checks["memory"] = {"status": "unknown", "message": "psutil not installed"}

    

    # 5. 系统信息

    checks["system"] = {

        "python_version": platform.python_version(),

        "platform": platform.system(),

        "version": "4.0.0",

        "indexing_mode": indexing_mode,

    }



    # 关键词模式下整体状态为 degraded（功能可用但非最佳）

    if indexing_mode == "keyword" and overall == "healthy":

        overall = "degraded"



    return {

        "status": overall,

        "checks": checks,

        "timestamp": time.time(),

    }





# ===== 智能体同步接口 =====



class AgentSyncItem(BaseModel):

    id: str

    name: str = ""

    task: str = ""

    mode: str = "agent"

    created_at: float = None

    updated_at: float = None



class AgentSyncRequest(BaseModel):

    agents: list = []



@router.post("/agents/sync", summary="同步智能体数据")

async def agents_sync(req: AgentSyncRequest, authorization: str = Header(None)):

    """

    同步智能体数据到服务端（按agent_id合并，updated_at较新的优先）

    用于跨浏览器/跨设备同步智能体prompt编辑

    """

    username = None

    if authorization and authorization.startswith("Bearer "):

        username = get_username_from_token(authorization[7:])

    if not username:

        raise HTTPException(status_code=401, detail="未登录")

    

    result = storage_sync_agents(username, req.agents)

    return {"success": True, "agents": result["agents"], "synced": result.get("synced", 0), "updated": result.get("updated", 0)}



@router.get("/agents", summary="获取用户智能体列表")

async def get_agents(authorization: str = Header(None)):

    """

    获取当前用户的智能体列表

    """

    username = None

    if authorization and authorization.startswith("Bearer "):

        username = get_username_from_token(authorization[7:])

    if not username:

        raise HTTPException(status_code=401, detail="未登录")

    

    agents = storage_load_agents(username)

    return {"success": True, "agents": agents}






# ===== 外部知识库上传 =====

@router.post("/external-kb/upload", summary="上传文档到全质知识库")
async def external_kb_upload(file: UploadFile = File(...), category: str = Form(""), subcategory: str = Form(""), subsubcategory: str = Form(""), admin: str = Depends(require_admin)):
    """上传文档到外部知识库（external_kb collection），按一级分类+二级+三级子目录存储"""
    allowed_ext = {".pdf", ".txt", ".md", ".docx", ".xlsx", ".xls", ".doc"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的格式: {ext}")

    # 保存文件到外部知识库目录（按一级分类 + 二级 + 三级子目录存子目录）
    ext_doc_dir = os.path.join(settings.DOCUMENTS_DIR, "external_kb")
    if category:
        ext_doc_dir = os.path.join(ext_doc_dir, category)
    if subcategory:
        ext_doc_dir = os.path.join(ext_doc_dir, subcategory)
    if subsubcategory:
        ext_doc_dir = os.path.join(ext_doc_dir, subsubcategory)
    os.makedirs(ext_doc_dir, exist_ok=True)
    decoded_filename = file.filename
    file_path = os.path.join(ext_doc_dir, decoded_filename)

    content_bytes = await file.read()
    with open(file_path, "wb") as f:
        f.write(content_bytes)

    logger.info(f"[外部知识库上传] 用户={username}, 文件={decoded_filename}, 分类={category}/{subcategory}/{subsubcategory}")

    # 索引到 external_kb collection（带 category + subcategory + subsubcategory 元数据）
    try:
        index_result = await asyncio.to_thread(index_document, file_path, decoded_filename,
                                                agent_id="__external__",
                                                category=category or None,
                                                subcategory=subcategory or None,
                                                subsubcategory=subsubcategory or None)
        if index_result.get("status") == "error":
            raise HTTPException(status_code=500, detail=index_result.get("message", "索引失败"))
        chunks = index_result.get("chunks", 0)
        logger.info(f"[外部知识库] 索引成功: {decoded_filename}, 分块={chunks}")
        return {"success": True, "filename": decoded_filename, "chunks": chunks, "message": "上传并索引成功"}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"[外部知识库] 索引异常: {e}")
        raise HTTPException(status_code=500, detail=f"索引失败: {str(e)}")


@router.get("/external-kb/documents", summary="列出全质知识库文档")
async def external_kb_documents(category: str = Query(None), subcategory: str = Query(None), subsubcategory: str = Query(None), username: str = Depends(require_auth)):
    """列出外部知识库的所有文档（按一级+二级+三级子目录过滤）"""
    try:
        docs = await asyncio.to_thread(list_indexed_documents, agent_id="__external__",
                                        category=category, subcategory=subcategory, subsubcategory=subsubcategory)
        return {"success": True, "documents": docs}
    except Exception as e:
        return {"success": True, "documents": []}


@router.get("/external-kb/stats", summary="全质知识库统计")
async def external_kb_stats(category: str = Query(None), username: str = Depends(require_auth)):
    """获取外部知识库的文档数和切片数"""
    try:
        docs = await asyncio.to_thread(list_indexed_documents, agent_id="__external__", category=category)
        doc_count = len(docs)
        # 从 ChromaDB 获取切片数
        chunk_count = 0
        try:
            from app.rag.document import get_vector_store
            vs = get_vector_store(agent_id="__external__")
            if vs is not None:
                collection = vs._collection
                all_docs = collection.get(include=["metadatas"])
                for meta in all_docs.get("metadatas", []):
                    if category:
                        if meta and meta.get("category") == category:
                            chunk_count += 1
                    else:
                        chunk_count += 1
        except Exception:
            pass
        return {"success": True, "doc_count": doc_count, "chunk_count": chunk_count}
    except Exception as e:
        return {"success": True, "doc_count": 0, "chunk_count": 0}


@router.delete("/external-kb/documents/{filename}", summary="删除全质知识库文档")
async def delete_external_kb_document(filename: str, admin: str = Depends(require_admin)):
    """删除外部知识库的文档（仅管理员可删除，普通用户只读+上传）"""
    try:
        result = await asyncio.to_thread(delete_document, filename, agent_id="__external__")
        return {"success": True, "message": f"已删除 {filename}"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ===== 一键生成质量手册（SCskill + AI 驱动）=====

# 把 SCskill scripts 目录加入 path，以便直接 import
_PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
_SCSCRIPTS_DIR = os.path.join(_PROJECT_ROOT, "skills", "SCskill", "scripts")
import sys as _sys_top
if _SCSCRIPTS_DIR not in _sys_top.path:
    _sys_top.path.insert(0, _SCSCRIPTS_DIR)


@router.post("/generate/manual", summary="一键生成质量手册（AI 驱动，SSE 流式进度）")
async def generate_manual_api(request: Request, username: str = Depends(require_auth)):
    """AI 驱动生成质量手册（支持多模板批量生成），SSE 流式推送每一步进度。
    用户可能在企业内部文件上传多个手册分册，每个都会生成一个文件。"""
    import asyncio
    import re as _re
    import sys as _sys
    import importlib
    import json as _json
    from datetime import datetime
    from langchain_core.messages import HumanMessage, SystemMessage
    from app.agent.core import create_llm

    body = await request.json()
    survey_data = body.get("survey_data", {})
    current_agent_id = body.get("agent_id", "")

    if not survey_data:
        raise HTTPException(status_code=400, detail="未提供体系调研数据")

    if not os.path.exists(os.path.join(_SCSCRIPTS_DIR, "generate_manual.py")):
        raise HTTPException(status_code=500, detail=f"SCskill 模块未找到: {_SCSCRIPTS_DIR}")

    export_dir = os.path.join(settings.DATA_DIR, "export")
    session_id_for_export = body.get("session_id", "")
    if session_id_for_export:
        export_dir = os.path.join(export_dir, session_id_for_export)
    os.makedirs(export_dir, exist_ok=True)

    # 优先使用前端传来的 model_id，确保与用户在下拉框的选择一致
    request_model_id = (body.get("model_id") or "").strip()
    valid_model_ids = set(m["id"] for m in AVAILABLE_MODELS)
    current_model = request_model_id if request_model_id in valid_model_ids else (get_current_model() or "glm-5.2")
    if request_model_id and request_model_id in valid_model_ids:
        logger.info(f"[SCskill] 使用前端指定模型: {current_model}")
    else:
        logger.info(f"[SCskill] 使用全局模型: {current_model}")

    async def sse_event_stream():
        async def send(evt: dict):
            return f"data: {_json.dumps(evt, ensure_ascii=False)}\n\n"

        try:
            yield await send({"type": "progress", "step": "分析调研数据", "message": f"正在分析体系调研数据（公司：{survey_data.get('sv_company_name','未填写')}）...", "progress": 5})
            await asyncio.sleep(0.1)

            if 'generate_manual' in _sys.modules:
                gm = importlib.reload(_sys.modules['generate_manual'])
            else:
                gm = importlib.import_module('generate_manual')

            yield await send({"type": "progress", "step": "查找模板", "message": "正在查找模板文件（企业内部文件→全质知识库）...", "progress": 10})
            await asyncio.sleep(0.1)
            user_agent_id = _user_agent_id(current_agent_id, username)
            all_templates = gm.find_all_templates(agent_id=user_agent_id, documents_dir=settings.DOCUMENTS_DIR)
            if not all_templates:
                yield await send({"type": "error", "message": "未找到模板文件。请在企业内部文件知识库[手册]分类或全质知识库[体系文件/手册/全质手册模板]下上传 .docx/.doc 模板。"})
                return

            total_templates = len(all_templates)
            yield await send({"type": "progress", "step": "已找到模板", "message": f"共找到 {total_templates} 个手册模板，将逐个生成", "progress": 12})
            for t in all_templates:
                source_label = {'internal': '企业内部文件', 'external': '全质知识库'}.get(t['source'], '未知')
                yield await send({"type": "progress", "step": "模板列表", "message": f"  [{source_label}] {t['filename']}", "progress": 12})

            survey_text = gm.format_survey_for_llm(survey_data)
            company_name = (survey_data.get('sv_company_name') or '企业').strip()
            safe_name = _re.sub(r'[\\/:*?"<>|]', '_', company_name)
            today_str = datetime.now().strftime("%Y%m%d")
            generated_files = []  # 主文件（AI 修改后的手册）
            reference_files = []  # 补缺参考文件（仅含用户手册未覆盖的章节）
            failed_files = []

            # [新需求] 拆分循环：先处理所有 internal 模板，再处理 external 模板
            # - internal 模板：AI 修改 → 主文件
            # - external 模板：
            #   * 若用户未上传任何手册 → 直接在 external 上 AI 修改 → 主文件（原流程）
            #   * 若用户已上传手册 → 提取差异章节 → 补缺参考文件
            internal_templates = [t for t in all_templates if t['source'] == 'internal']
            external_templates = [t for t in all_templates if t['source'] == 'external']
            total_templates = len(all_templates)
            ti_counter = 0  # 全局计数器，用于进度计算
            # 收集所有用户手册的章节文本，后续用于与 external 对比
            user_manual_sections_text = ''

            # ========== 阶段 1：处理 internal 模板 ==========
            for tmpl in internal_templates:
                ti = ti_counter
                ti_counter += 1
                template_path = tmpl['path']
                need_convert = tmpl['need_convert']
                template_source = tmpl['source']
                template_filename = tmpl['filename']
                source_label = {'internal': '企业内部文件', 'external': '全质知识库'}.get(template_source, '未知')
                base_progress = int(15 + (ti / total_templates) * 80)
                end_progress = int(15 + ((ti + 1) / total_templates) * 80)

                yield await send({"type": "progress", "step": f"处理第 {ti+1}/{total_templates} 个：{template_filename}", "message": f"正在处理 [{source_label}] {template_filename}", "progress": base_progress})

                try:
                    actual_template = template_path
                    if need_convert:
                        converted = await asyncio.to_thread(gm.convert_doc_to_docx, template_path)
                        if not converted:
                            failed_files.append({"filename": template_filename, "reason": "doc转换失败"})
                            continue
                        actual_template = converted

                    # [主文件] internal 模板调 AI 修改
                    from docx import Document
                    doc = await asyncio.to_thread(Document, str(actual_template))
                    overview = await asyncio.to_thread(gm.extract_template_overview, doc)
                    overview_text = gm.format_overview_for_llm(overview)
                    # 累积用户手册章节文本（用于后续与 external 对比）
                    user_manual_sections_text += f"\n=== 用户手册：{template_filename} ===\n" + overview_text + "\n"
                    system_prompt, user_prompt = gm.build_llm_prompt(overview_text, survey_text, survey_data)

                    yield await send({"type": "progress", "step": f"AI分析 {template_filename}", "message": f"正在调用 AI 分析...", "progress": base_progress + 10})

                    llm = create_llm(deep_think=True, model_override=current_model)
                    messages = [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]
                    stats = {'paragraph': 0, 'table_cell': 0, 'global_replace': 0, 'header_replace': 0, 'unknown': 0, 'failed': 0}
                    modifications_count = 0
                    buffer = ''
                    try:
                        async for chunk in llm.astream(messages):
                            if not chunk: continue
                            token = chunk.content if hasattr(chunk, 'content') and isinstance(chunk.content, str) else str(chunk)
                            if not token: continue
                            buffer += token
                            while '\n' in buffer:
                                line, buffer = buffer.split('\n', 1)
                                line = line.strip()
                                if not line or line == '===END===': continue
                                if line.startswith('```'): line = line.lstrip('`').replace('json', '', 1).strip()
                                if line.endswith('```'): line = line[:-3].strip()
                                if not line.startswith('{') or not line.endswith('}'): continue
                                try: mod = _json.loads(line)
                                except: 
                                    try: mod = _json.loads(_re.sub(r',\s*([}\]])', r'\1', line))
                                    except: continue
                                modifications_count += 1
                                mod_type = mod.get('type', '')
                                reason = (mod.get('reason', '') or '')[:80]
                                progress = int(base_progress + 10 + (modifications_count % 20) * 2)
                                try:
                                    if mod_type == 'paragraph':
                                        idx = int(mod.get('index', -1)); new_text = mod.get('new_text', '')
                                        if gm.apply_paragraph_replace(doc, idx, new_text): stats['paragraph'] += 1
                                        else: stats['failed'] += 1
                                        yield await send({"type": "modification", "mod_type": "paragraph", "location": f"P{idx}", "reason": reason, "preview": (new_text[:60] + '...' if len(new_text) > 60 else new_text), "progress": progress, "index": modifications_count})
                                    elif mod_type == 'table_cell':
                                        ti2 = int(mod.get('table', -1)); ri = int(mod.get('row', -1)); ci2 = int(mod.get('col', -1)); new_text = mod.get('new_text', '')
                                        if gm.apply_table_cell_replace(doc, ti2, ri, ci2, new_text): stats['table_cell'] += 1
                                        else: stats['failed'] += 1
                                        yield await send({"type": "modification", "mod_type": "table_cell", "location": f"T{ti2}.R{ri}.C{ci2}", "reason": reason, "preview": (new_text[:60] + '...' if len(new_text) > 60 else new_text), "progress": progress, "index": modifications_count})
                                    elif mod_type == 'global_replace':
                                        old = mod.get('old', ''); new = mod.get('new', ''); n = gm.apply_global_replace(doc, old, new)
                                        stats['global_replace'] += n
                                        yield await send({"type": "modification", "mod_type": "global_replace", "location": "全文", "reason": reason, "preview": f"'{old}' → '{new}'（{n} 处）", "progress": progress, "index": modifications_count})
                                    elif mod_type == 'header_replace':
                                        old = mod.get('old', ''); new = mod.get('new', ''); n = gm.apply_header_replace(doc, old, new)
                                        stats['header_replace'] += n
                                        yield await send({"type": "modification", "mod_type": "header_replace", "location": "页眉/页脚", "reason": reason, "preview": f"'{old}' → '{new}'（{n} 处）", "progress": progress, "index": modifications_count})
                                    else: stats['unknown'] += 1
                                except: stats['failed'] += 1
                    except Exception as stream_err:
                        logger.warning(f"[SCskill] {template_filename} LLM异常: {stream_err}")

                    if modifications_count == 0:
                        failed_files.append({"filename": template_filename, "reason": "AI未生成修改方案"})
                        continue

                    tmpl_base = os.path.splitext(template_filename)[0]
                    tmpl_clean = _re.sub(r'^\d*-?[A-Z]+-\w+-\w+-?\d*\s*', '', tmpl_base)
                    tmpl_clean = _re.sub(r'^[A-Z]+-\w+-\w+-?\d*\s*', '', tmpl_clean)
                    tmpl_clean = _re.sub(r'-A\d+$', '', tmpl_clean)
                    tmpl_clean = tmpl_clean.strip() or '质量管理手册'
                    safe_tmpl_name = _re.sub(r'[\\/:*?"<>|]', '_', tmpl_clean)
                    out_filename = f"{safe_tmpl_name}_{safe_name}_{today_str}.docx"
                    output_path = os.path.join(export_dir, out_filename)
                    gm.remove_even_page_headers_footers(doc)
                    await asyncio.to_thread(doc.save, output_path)
                    # 清理 doc 转换产生的临时目录
                    if need_convert and converted:
                        try:
                            import shutil
                            shutil.rmtree(os.path.dirname(converted), ignore_errors=True)
                        except Exception:
                            pass
                    logger.info(f"[SCskill] 手册已生成: {output_path}")

                    download_url = f"/api/v1/documents/export-download/{out_filename}?sid={session_id_for_export}"
                    generated_files.append({"filename": out_filename, "download_url": download_url, "display_name": tmpl_clean, "modifications_count": modifications_count, "stats": stats})
                    yield await send({"type": "file_complete", "filename": out_filename, "download_url": download_url, "display_name": tmpl_clean, "modifications_count": modifications_count, "template_source": source_label, "template_filename": template_filename, "progress": end_progress, "is_reference": False})

                except Exception as e:
                    logger.exception(f"[SCskill] {template_filename} 生成异常: {e}")
                    failed_files.append({"filename": template_filename, "reason": str(e)[:100]})

            # ========== 阶段 2：处理 external 模板 ==========
            for tmpl in external_templates:
                ti = ti_counter
                ti_counter += 1
                template_path = tmpl['path']
                need_convert = tmpl['need_convert']
                template_source = tmpl['source']
                template_filename = tmpl['filename']
                source_label = {'internal': '企业内部文件', 'external': '全质知识库'}.get(template_source, '未知')
                base_progress = int(15 + (ti / total_templates) * 80)
                end_progress = int(15 + ((ti + 1) / total_templates) * 80)

                yield await send({"type": "progress", "step": f"处理第 {ti+1}/{total_templates} 个：{template_filename}", "message": f"正在处理 [{source_label}] {template_filename}", "progress": base_progress})

                try:
                    actual_template = template_path
                    if need_convert:
                        converted = await asyncio.to_thread(gm.convert_doc_to_docx, template_path)
                        if not converted:
                            failed_files.append({"filename": template_filename, "reason": "doc转换失败"})
                            continue
                        actual_template = converted

                    # 分支 A：用户未上传任何手册 → 直接在 external 上 AI 修改 → 主文件（原流程）
                    if not internal_templates or not user_manual_sections_text.strip():
                        yield await send({"type": "progress", "step": f"AI分析 {template_filename}", "message": f"用户未上传手册，在全质模板上修改...", "progress": base_progress + 5})
                        from docx import Document
                        doc = await asyncio.to_thread(Document, str(actual_template))
                        overview = await asyncio.to_thread(gm.extract_template_overview, doc)
                        overview_text = gm.format_overview_for_llm(overview)
                        system_prompt, user_prompt = gm.build_llm_prompt(overview_text, survey_text, survey_data)
                        llm = create_llm(deep_think=True, model_override=current_model)
                        messages = [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]
                        stats = {'paragraph': 0, 'table_cell': 0, 'global_replace': 0, 'header_replace': 0, 'unknown': 0, 'failed': 0}
                        modifications_count = 0
                        buffer = ''
                        try:
                            async for chunk in llm.astream(messages):
                                if not chunk: continue
                                token = chunk.content if hasattr(chunk, 'content') and isinstance(chunk.content, str) else str(chunk)
                                if not token: continue
                                buffer += token
                                while '\n' in buffer:
                                    line, buffer = buffer.split('\n', 1)
                                    line = line.strip()
                                    if not line or line == '===END===': continue
                                    if line.startswith('```'): line = line.lstrip('`').replace('json', '', 1).strip()
                                    if line.endswith('```'): line = line[:-3].strip()
                                    if not line.startswith('{') or not line.endswith('}'): continue
                                    try: mod = _json.loads(line)
                                    except:
                                        try: mod = _json.loads(_re.sub(r',\s*([}\]])', r'\1', line))
                                        except: continue
                                    modifications_count += 1
                                    mod_type = mod.get('type', '')
                                    reason = (mod.get('reason', '') or '')[:80]
                                    progress = int(base_progress + 10 + (modifications_count % 20) * 2)
                                    try:
                                        if mod_type == 'paragraph':
                                            idx = int(mod.get('index', -1)); new_text = mod.get('new_text', '')
                                            if gm.apply_paragraph_replace(doc, idx, new_text): stats['paragraph'] += 1
                                            else: stats['failed'] += 1
                                            yield await send({"type": "modification", "mod_type": "paragraph", "location": f"P{idx}", "reason": reason, "preview": (new_text[:60] + '...' if len(new_text) > 60 else new_text), "progress": progress, "index": modifications_count})
                                        elif mod_type == 'table_cell':
                                            ti2 = int(mod.get('table', -1)); ri = int(mod.get('row', -1)); ci2 = int(mod.get('col', -1)); new_text = mod.get('new_text', '')
                                            if gm.apply_table_cell_replace(doc, ti2, ri, ci2, new_text): stats['table_cell'] += 1
                                            else: stats['failed'] += 1
                                            yield await send({"type": "modification", "mod_type": "table_cell", "location": f"T{ti2}.R{ri}.C{ci2}", "reason": reason, "preview": (new_text[:60] + '...' if len(new_text) > 60 else new_text), "progress": progress, "index": modifications_count})
                                        elif mod_type == 'global_replace':
                                            old = mod.get('old', ''); new = mod.get('new', ''); n = gm.apply_global_replace(doc, old, new)
                                            stats['global_replace'] += n
                                            yield await send({"type": "modification", "mod_type": "global_replace", "location": "全文", "reason": reason, "preview": f"'{old}' → '{new}'（{n} 处）", "progress": progress, "index": modifications_count})
                                        elif mod_type == 'header_replace':
                                            old = mod.get('old', ''); new = mod.get('new', ''); n = gm.apply_header_replace(doc, old, new)
                                            stats['header_replace'] += n
                                            yield await send({"type": "modification", "mod_type": "header_replace", "location": "页眉/页脚", "reason": reason, "preview": f"'{old}' → '{new}'（{n} 处）", "progress": progress, "index": modifications_count})
                                        else: stats['unknown'] += 1
                                    except: stats['failed'] += 1
                        except Exception as stream_err:
                            logger.warning(f"[SCskill] {template_filename} LLM异常: {stream_err}")

                        if modifications_count == 0:
                            failed_files.append({"filename": template_filename, "reason": "AI未生成修改方案"})
                            continue

                        tmpl_base = os.path.splitext(template_filename)[0]
                        tmpl_clean = _re.sub(r'^\d*-?[A-Z]+-\w+-\w+-?\d*\s*', '', tmpl_base)
                        tmpl_clean = _re.sub(r'^[A-Z]+-\w+-\w+-?\d*\s*', '', tmpl_clean)
                        tmpl_clean = _re.sub(r'-A\d+$', '', tmpl_clean)
                        tmpl_clean = tmpl_clean.strip() or '质量管理手册'
                        safe_tmpl_name = _re.sub(r'[\\/:*?"<>|]', '_', tmpl_clean)
                        out_filename = f"{safe_tmpl_name}_{safe_name}_{today_str}.docx"
                        output_path = os.path.join(export_dir, out_filename)
                        gm.remove_even_page_headers_footers(doc)
                        await asyncio.to_thread(doc.save, output_path)
                        if need_convert and converted:
                            try:
                                import shutil
                                shutil.rmtree(os.path.dirname(converted), ignore_errors=True)
                            except Exception:
                                pass
                        logger.info(f"[SCskill] 手册已生成(全质模板上修改): {output_path}")
                        download_url = f"/api/v1/documents/export-download/{out_filename}?sid={session_id_for_export}"
                        generated_files.append({"filename": out_filename, "download_url": download_url, "display_name": tmpl_clean, "modifications_count": modifications_count, "stats": stats})
                        yield await send({"type": "file_complete", "filename": out_filename, "download_url": download_url, "display_name": tmpl_clean, "modifications_count": modifications_count, "template_source": source_label, "template_filename": template_filename, "progress": end_progress, "is_reference": False})
                        continue

                    # 分支 B：用户已上传手册 → 提取差异章节 → 补缺参考文件
                    yield await send({"type": "progress", "step": f"对比分析 {template_filename}", "message": f"正在对比用户手册与全质模板，提取差异章节...", "progress": base_progress + 5})
                    from docx import Document as _DocExt
                    ext_doc = await asyncio.to_thread(_DocExt, str(actual_template))
                    ext_overview = await asyncio.to_thread(gm.extract_template_overview, ext_doc)
                    ext_overview_text = gm.format_overview_for_llm(ext_overview)

                    # 调 AI 提取差异章节
                    diff_system, diff_user = gm.build_diff_prompt(ext_overview_text, user_manual_sections_text, template_filename)
                    diff_llm = create_llm(deep_think=True, model_override=current_model)
                    diff_messages = [SystemMessage(content=diff_system), HumanMessage(content=diff_user)]
                    diff_buffer = ''
                    diff_sections = []  # 收集差异章节文本
                    try:
                        async for chunk in diff_llm.astream(diff_messages):
                            if not chunk: continue
                            token = chunk.content if hasattr(chunk, 'content') and isinstance(chunk.content, str) else str(chunk)
                            if not token: continue
                            diff_buffer += token
                            while '\n' in diff_buffer:
                                line, diff_buffer = diff_buffer.split('\n', 1)
                                line = line.strip()
                                if not line or line == '===END===': continue
                                if line.startswith('```'): line = line.lstrip('`').replace('json', '', 1).strip()
                                if line.endswith('```'): line = line[:-3].strip()
                                if line.startswith('差异章节:') or line.startswith('差异章节：'):
                                    # 文本格式：差异章节: 标题 - 正文
                                    diff_sections.append(line)
                                    yield await send({"type": "progress", "step": f"提取差异", "message": line[:80], "progress": int(base_progress + 20 + len(diff_sections) * 5)})
                    except Exception as stream_err:
                        logger.warning(f"[SCskill] {template_filename} 差异提取异常: {stream_err}")

                    # 清理 doc 转换产生的临时目录
                    if need_convert and converted:
                        try:
                            import shutil
                            shutil.rmtree(os.path.dirname(converted), ignore_errors=True)
                        except Exception:
                            pass

                    if not diff_sections:
                        yield await send({"type": "progress", "step": f"无差异 {template_filename}", "message": f"用户手册已涵盖全质模板内容，无需生成参考文件", "progress": end_progress})
                        logger.info(f"[SCskill] {template_filename}: 无差异章节，跳过参考文件")
                        continue

                    # 生成补缺参考文件 docx
                    tmpl_base = os.path.splitext(template_filename)[0]
                    tmpl_clean = _re.sub(r'^\d*-?[A-Z]+-\w+-\w+-?\d*\s*', '', tmpl_base)
                    tmpl_clean = _re.sub(r'^[A-Z]+-\w+-\w+-?\d*\s*', '', tmpl_clean)
                    tmpl_clean = _re.sub(r'-A\d+$', '', tmpl_clean)
                    tmpl_clean = tmpl_clean.strip() or '质量管理手册'
                    safe_tmpl_name = _re.sub(r'[\\/:*?"<>|]', '_', tmpl_clean)
                    ref_filename = f"补缺参考_{safe_tmpl_name}_{safe_name}_{today_str}.docx"
                    ref_output_path = os.path.join(export_dir, ref_filename)
                    # 生成 docx：标题 + 差异章节文本
                    from docx import Document as _DocOut
                    from docx.shared import Pt
                    ref_doc = _DocOut()
                    # 标题
                    title_p = ref_doc.add_paragraph()
                    title_run = title_p.add_run(f"补缺参考文件 — {tmpl_clean}")
                    title_run.bold = True
                    title_run.font.size = Pt(16)
                    # 说明
                    note_p = ref_doc.add_paragraph()
                    note_run = note_p.add_run(f"本文档为「全质知识库模板中包含、但用户上传手册中未覆盖」的章节内容，供补缺参考。\n来源模板：{template_filename}\n生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
                    note_run.font.size = Pt(10)
                    ref_doc.add_paragraph('')  # 空行
                    # 逐个差异章节
                    for sec in diff_sections:
                        p = ref_doc.add_paragraph(sec)
                        ref_doc.add_paragraph('')  # 章节间空行
                    await asyncio.to_thread(ref_doc.save, ref_output_path)
                    logger.info(f"[SCskill] 补缺参考文件已生成: {ref_output_path}")
                    ref_download_url = f"/api/v1/documents/export-download/{ref_filename}?sid={session_id_for_export}"
                    reference_files.append({"filename": ref_filename, "download_url": ref_download_url, "display_name": tmpl_clean + "（补缺参考）", "template_source": source_label, "template_filename": template_filename})
                    yield await send({"type": "file_complete", "filename": ref_filename, "download_url": ref_download_url, "display_name": tmpl_clean + "（补缺参考）", "modifications_count": len(diff_sections), "template_source": source_label + "（补缺参考）", "template_filename": template_filename, "progress": end_progress, "is_reference": True})

                except Exception as e:
                    logger.exception(f"[SCskill] {template_filename} 生成异常: {e}")
                    failed_files.append({"filename": template_filename, "reason": str(e)[:100]})

            # 保存对话记录
            try:
                from app.memory.manager import get_session_history, flush_session
                from langchain_core.messages import AIMessage as _AIMsg2
                history = get_session_history(session_id_for_export)
                if generated_files or reference_files:
                    parts = []
                    if generated_files:
                        file_list = "\n".join([f"- {f.get('display_name', f.get('filename',''))}（{f.get('modifications_count',0)} 处修改）" for f in generated_files])
                        parts.append(f"已生成 {len(generated_files)} 个质量手册：\n{file_list}")
                    if reference_files:
                        ref_list = "\n".join([f"- {f.get('display_name', f.get('filename',''))}（补缺内容，非全质模板副本）" for f in reference_files])
                        parts.append(f"另附 {len(reference_files)} 个补缺参考文件（全质模板中包含但用户手册未覆盖的章节）：\n{ref_list}")
                    ai_msg = "\n\n".join(parts)
                    # DOWNLOAD_INFO 只含主文件，参考文件用单独的 REFERENCE_INFO 标记
                    download_info = {"type": "manual_files", "files": [{"filename": f.get("filename",""), "download_url": f.get("download_url",""), "display_name": f.get("display_name", f.get("filename","")), "modifications_count": f.get("modifications_count", 0)} for f in generated_files]}
                    reference_info = {"type": "reference_files", "files": [{"filename": f.get("filename",""), "download_url": f.get("download_url",""), "display_name": f.get("display_name", f.get("filename",""))} for f in reference_files]}
                    ai_msg = f"<!--DOWNLOAD_INFO:{_json.dumps(download_info, ensure_ascii=False)}--><!--REFERENCE_INFO:{_json.dumps(reference_info, ensure_ascii=False)}-->\n\n" + ai_msg
                else:
                    ai_msg = "生成失败，未产生任何手册文件。"
                history.add_message(_AIMsg2(content=ai_msg))
                parts = session_id_for_export.split("_", 1)
                if len(parts) == 2: update_chat_time(parts[0], session_id_for_export)
                flush_session(session_id_for_export)
            except Exception as save_err:
                logger.warning(f"[SCskill] 保存对话记录失败: {save_err}")

            yield await send({"type": "success", "total_files": len(generated_files), "reference_count": len(reference_files), "failed_count": len(failed_files), "files": generated_files, "reference_files": reference_files, "failed_files": failed_files, "model_used": current_model, "progress": 100})

        except Exception as e:
            logger.exception(f"[SCskill] 生成手册异常: {e}")
            yield await send({"type": "error", "message": f"生成失败: {str(e)}"})

    return StreamingResponse(sse_event_stream(), media_type="text/event-stream", headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"})



# ===== 一键生成程序文件（CXskill + AI 驱动）=====

_CXSCRIPTS_DIR = os.path.join(_PROJECT_ROOT, "skills", "CXskill", "scripts")
if _CXSCRIPTS_DIR not in _sys_top.path:
    _sys_top.path.insert(0, _CXSCRIPTS_DIR)


@router.post("/generate/procedure/templates", summary="列出可用的程序文件模板")
async def list_procedure_templates(request: Request, username: str = Depends(require_auth)):
    """列出所有可用的程序文件模板，供前端勾选。
    如果用户上传了文件（企业内部文件有模板），直接返回这些模板并标记 auto_generate=True。
    如果用户没上传（企业内部文件为空），返回全质知识库的模板列表供用户勾选。
    
    AI 智能去重：对每个部门，如果企业内部文件和全质知识库文件是"同一个程序"
    （即使文件名不同但内容主题相同，如"文件管理规则"="文件控制程序"），
    则跳过全质知识库的对应文件，避免重复生成。"""
    import importlib
    import sys as _sys
    import re as _re
    from langchain_core.messages import HumanMessage, SystemMessage
    from app.agent.core import create_llm

    body = await request.json()
    agent_id = body.get("agent_id", "")

    if 'generate_procedure' in _sys.modules:
        cx = importlib.reload(_sys.modules['generate_procedure'])
    else:
        cx = importlib.import_module('generate_procedure')

    all_templates = cx.find_all_templates(agent_id=_user_agent_id(agent_id, username), documents_dir=settings.DOCUMENTS_DIR)

    # 检查企业内部文件是否有模板
    internal_count = sum(1 for t in all_templates if t['source'] == 'internal')
    external_count = sum(1 for t in all_templates if t['source'] == 'external')

    # ===== AI 智能去重 =====
    # 对每个部门，找出企业内部文件和全质知识库文件中"内容主题相同"的，跳过全质知识库的对应文件
    # 仅当企业内部和全质知识库都存在文件时才需要去重；企业内部为空时无需去重
    templates_to_skip = set()  # 存放需要跳过的 (dept, filename, source=external) 元组
    if internal_count > 0 and external_count > 0:
        # 按部门分组
        from collections import defaultdict
        dept_internal = defaultdict(list)  # {dept: [filename, ...]}
        dept_external = defaultdict(list)
        for t in all_templates:
            if t['source'] == 'internal':
                dept_internal[t['dept']].append(t['filename'])
            else:
                dept_external[t['dept']].append(t['filename'])

        # 只对"企业内部和全质知识库都有文件"的部门做 AI 去重
        try:
            llm = create_llm(short_response=True)
            for dept in dept_internal:
                internal_files = dept_internal[dept]
                external_files = dept_external.get(dept, [])
                if not external_files:
                    continue  # 全质知识库没文件，无需去重
                # 如果企业内部文件数 >= 全质知识库文件数，且文件名相似度高，可能用户已经全量上传，可整体跳过
                # 但保险起见还是逐个让 AI 判断
                
                dedup_prompt = f"""你是企业质量管理文件专家。请判断以下两组文件是否包含"同一个程序文件"（内容主题相同，只是文件名表述不同）。

部门：{dept}

企业内部上传的文件：
{chr(10).join(f"{i+1}. {f}" for i, f in enumerate(internal_files))}

全质知识库的文件：
{chr(10).join(f"{i+1}. {f}" for i, f in enumerate(external_files))}

判断规则：
- "文件控制程序"、"文件管理规则"、"文件管理制度" → 是同一个程序
- "记录控制程序"、"记录管理程序"、"记录管理规则" → 是同一个程序
- "内部审核程序"、"内部审核控制程序"、"内审程序" → 是同一个程序
- 即：核心主题词（文件控制/记录控制/内部审核/管理评审/不合格品控制/纠正措施/预防措施等）相同的，就是同一个程序
- 即使文件名格式完全不同（如带编号"4-AAA-QA-QP-01 xxx"vs不带编号"xxx"），只要主题相同就算同一个

请严格按以下JSON格式输出，只输出JSON：
{{"skip_external":[需要跳过的全质知识库文件名列表，即被企业内部文件覆盖的]}}
不确定的不要放入skip_external（宁可保留生成）。"""

                messages = [
                    SystemMessage(content="你是企业质量管理文件专家，只输出JSON格式结果。"),
                    HumanMessage(content=dedup_prompt)
                ]
                response = await llm.ainvoke(messages)
                ai_text = response.content.strip()
                json_match = _re.search(r'\{[\s\S]*\}', ai_text)
                if json_match:
                    try:
                        import json as _json
                        result = _json.loads(json_match.group(0))
                        skip_list = result.get('skip_external', [])
                        for skip_file in skip_list:
                            # 确保是字符串且非空
                            if isinstance(skip_file, str) and skip_file.strip():
                                templates_to_skip.add((dept, skip_file.strip(), 'external'))
                                logger.info(f"[程序模板AI去重] 部门={dept}, 跳过全质文件={skip_file.strip()}（被企业内部覆盖）")
                    except Exception as parse_err:
                        logger.warning(f"[程序模板AI去重] 解析失败 dept={dept}: {parse_err}, raw={ai_text[:200]}")
        except Exception as e:
            logger.warning(f"[程序模板AI去重] 整体失败，跳过去重（保留所有模板）: {e}")

    # 过滤掉被去重的模板
    filtered_templates = []
    for t in all_templates:
        if (t['dept'], t['filename'], t['source']) in templates_to_skip:
            continue
        filtered_templates.append(t)

    templates_list = []
    for t in filtered_templates:
        source_label = {'internal': '企业内部文件', 'external': '全质知识库'}.get(t['source'], '未知')
        templates_list.append({
            "dept": t['dept'],
            "filename": t['filename'],
            "source": t['source'],
            "source_label": source_label,
            "need_convert": t['need_convert'],
        })

    return {
        "success": True,
        "templates": templates_list,
        "auto_generate": internal_count > 0,  # 用户上传了就自动生成，不需要勾选
        "internal_count": internal_count,
        "external_count": sum(1 for t in all_templates if t['source'] == 'external'),
    }


@router.post("/generate/procedure", summary="一键生成程序文件（AI 驱动，SSE 流式进度，多文件批量生成）")
async def generate_procedure_api(request: Request, username: str = Depends(require_auth)):
    """AI 驱动生成程序文件（批量），SSE 流式推送每一步进度。
    查找所有部门的模板，逐个 AI 修改，返回多个下载链接。"""
    import asyncio
    import re as _re
    import sys as _sys
    import importlib
    import json as _json
    from datetime import datetime
    from langchain_core.messages import HumanMessage, SystemMessage
    from app.agent.core import create_llm

    body = await request.json()
    survey_data = body.get("survey_data", {})
    current_agent_id = body.get("agent_id", "")
    session_id_for_export = body.get("session_id", "")
    selected_templates = body.get("selected_templates", None)  # 用户勾选的模板列表，None=全部

    if not survey_data:
        raise HTTPException(status_code=400, detail="未提供体系调研数据")

    if not os.path.exists(os.path.join(_CXSCRIPTS_DIR, "generate_procedure.py")):
        raise HTTPException(status_code=500, detail=f"CXskill 模块未找到: {_CXSCRIPTS_DIR}")

    export_dir = os.path.join(settings.DATA_DIR, "export")
    if session_id_for_export:
        export_dir = os.path.join(export_dir, session_id_for_export)
    os.makedirs(export_dir, exist_ok=True)

    # 优先使用前端传来的 model_id，确保与用户在下拉框的选择一致
    request_model_id = (body.get("model_id") or "").strip()
    valid_model_ids = set(m["id"] for m in AVAILABLE_MODELS)
    current_model = request_model_id if request_model_id in valid_model_ids else (get_current_model() or "glm-5.2")
    if request_model_id and request_model_id in valid_model_ids:
        logger.info(f"[CXskill] 使用前端指定模型: {current_model}")
    else:
        logger.info(f"[CXskill] 使用全局模型: {current_model}")

    async def sse_event_stream():
        async def send(evt: dict):
            return f"data: {_json.dumps(evt, ensure_ascii=False)}\n\n"

        try:
            # 步骤 1：分析调研数据
            yield await send({
                "type": "progress", "step": "分析调研数据",
                "message": f"正在分析体系调研数据（公司：{survey_data.get('sv_company_name','未填写')}）...",
                "progress": 5
            })
            await asyncio.sleep(0.1)

            # 步骤 2：import CXskill
            if 'generate_procedure' in _sys.modules:
                cx = importlib.reload(_sys.modules['generate_procedure'])
            else:
                cx = importlib.import_module('generate_procedure')

            # 步骤 3：查找所有模板
            yield await send({
                "type": "progress", "step": "查找模板",
                "message": "正在查找所有部门的程序文件模板...",
                "progress": 10
            })
            await asyncio.sleep(0.1)
            all_templates = cx.find_all_templates(
                agent_id=_user_agent_id(current_agent_id, username), documents_dir=settings.DOCUMENTS_DIR)

            # [Bug 修复] AI 智能去重：如果用户上传了文件，跳过全质知识库中同主题的文件
            # 避免用户上传的文件和全质知识库的文件重复生成
            internal_count = sum(1 for t in all_templates if t['source'] == 'internal')
            external_count = sum(1 for t in all_templates if t['source'] == 'external')
            if internal_count > 0 and external_count > 0:
                from collections import defaultdict as _dd
                dept_internal = _dd(list)
                dept_external = _dd(list)
                for t in all_templates:
                    if t['source'] == 'internal':
                        dept_internal[t['dept']].append(t['filename'])
                    else:
                        dept_external[t['dept']].append(t['filename'])
                templates_to_skip = set()
                try:
                    from langchain_core.messages import HumanMessage as _HM, SystemMessage as _SM
                    _dedup_llm = create_llm(short_response=True)
                    for dept in dept_internal:
                        internal_files = dept_internal[dept]
                        external_files = dept_external.get(dept, [])
                        if not external_files:
                            continue
                        dedup_prompt = f"""你是企业质量管理文件专家。请判断以下两组文件是否包含"同一个程序文件"（内容主题相同，只是文件名表述不同）。

部门：{dept}

企业内部上传的文件：
{chr(10).join(f"{i+1}. {f}" for i, f in enumerate(internal_files))}

全质知识库的文件：
{chr(10).join(f"{i+1}. {f}" for i, f in enumerate(external_files))}

判断规则：
- "文件控制程序"、"文件管理规则"、"文件管理制度" → 是同一个程序
- "记录控制程序"、"记录管理程序"、"记录管理规则" → 是同一个程序
- 即：核心主题词相同的，就是同一个程序
- 即使文件名格式完全不同，只要主题相同就算同一个

请严格按以下JSON格式输出，只输出JSON：
{{"skip_external":[需要跳过的全质知识库文件名列表，即被企业内部文件覆盖的]}}
不确定的不要放入skip_external（宁可保留生成）。"""
                        _dedup_messages = [_SM(content="你是企业质量管理文件专家，只输出JSON格式结果。"), _HM(content=dedup_prompt)]
                        _dedup_resp = await _dedup_llm.ainvoke(_dedup_messages)
                        _dedup_text = _dedup_resp.content.strip()
                        _json_m = _re.search(r'\{[\s\S]*\}', _dedup_text)
                        if _json_m:
                            try:
                                import json as _json_dedup
                                _dedup_result = _json_dedup.loads(_json_m.group(0))
                                _skip_list = _dedup_result.get('skip_external', [])
                                for _skip_file in _skip_list:
                                    if isinstance(_skip_file, str) and _skip_file.strip():
                                        templates_to_skip.add((dept, _skip_file.strip(), 'external'))
                                        logger.info(f"[程序生成AI去重] 部门={dept}, 跳过全质文件={_skip_file.strip()}")
                            except Exception as _pe:
                                logger.warning(f"[程序生成AI去重] 解析失败 dept={dept}: {_pe}")
                except Exception as _de:
                    logger.warning(f"[程序生成AI去重] 整体失败，跳过去重: {_de}")
                # 过滤
                if templates_to_skip:
                    before_count = len(all_templates)
                    all_templates = [t for t in all_templates if (t['dept'], t['filename'], t['source']) not in templates_to_skip]
                    logger.info(f"[程序生成AI去重] 去重前 {before_count} 个模板，去重后 {len(all_templates)} 个")

            # 如果用户勾选了模板，只保留勾选的
            if selected_templates and isinstance(selected_templates, list) and len(selected_templates) > 0:
                all_templates = [t for t in all_templates
                                 if f"{t['dept']}/{t['filename']}" in selected_templates]

            if not all_templates:
                yield await send({
                    "type": "error",
                    "message": "未找到任何程序文件模板。请在企业内部文件知识库[程序文件]分类或全质知识库[体系文件/程序文件]下上传 .docx/.doc 模板。"
                })
                return

            total_templates = len(all_templates)
            yield await send({
                "type": "progress", "step": "已找到模板",
                "message": f"共找到 {total_templates} 个模板，将逐个生成程序文件",
                "progress": 12
            })
            # 列出所有找到的模板
            for t in all_templates:
                source_label = {'internal': '企业内部文件', 'external': '全质知识库'}.get(t['source'], '未知')
                yield await send({
                    "type": "progress", "step": "模板列表",
                    "message": f"  [{source_label}] {t['dept']}/{t['filename']}",
                    "progress": 12
                })

            survey_text = cx.format_survey_for_llm(survey_data)
            company_name = (survey_data.get('sv_company_name') or '企业').strip()
            safe_name = _re.sub(r'[\\/:*?"<>|]', '_', company_name)
            today_str = datetime.now().strftime("%Y%m%d")

            generated_files = []
            failed_files = []

            # 逐个处理每个模板
            for ti, tmpl in enumerate(all_templates):
                dept = tmpl['dept']
                filename_tmpl = tmpl['filename']
                source = tmpl['source']
                source_label = {'internal': '企业内部文件知识库', 'external': '全质知识库'}.get(source, '未知')
                base_progress = int(15 + (ti / total_templates) * 80)  # 15% ~ 95%
                end_progress = int(15 + ((ti + 1) / total_templates) * 80)

                yield await send({
                    "type": "progress", "step": f"处理第 {ti+1}/{total_templates} 个：{dept}",
                    "message": f"正在处理 [{source_label}] {dept}/{filename_tmpl}",
                    "progress": base_progress
                })

                try:
                    # .doc 转 .docx
                    actual_template = tmpl['path']
                    if tmpl['need_convert']:
                        converted = await asyncio.to_thread(cx.convert_doc_to_docx, tmpl['path'])
                        if not converted:
                            yield await send({
                                "type": "progress", "step": f"跳过 {dept}",
                                "message": f"⚠️ {dept}/{filename_tmpl} .doc转换失败，跳过",
                                "progress": end_progress
                            })
                            failed_files.append({"dept": dept, "filename": filename_tmpl, "reason": "doc转换失败"})
                            continue
                        actual_template = converted

                    # 加载模板
                    from docx import Document
                    doc = await asyncio.to_thread(Document, str(actual_template))

                    # 提取概览
                    overview = await asyncio.to_thread(cx.extract_template_overview, doc)
                    overview_text = cx.format_overview_for_llm(overview)

                    # 构造提示词（传入 survey_data 用于计算实施日期）
                    system_prompt, user_prompt = cx.build_llm_prompt(overview_text, survey_text, filename_tmpl, survey_data)

                    # AI 修改
                    llm = create_llm(deep_think=True, model_override=current_model)
                    messages = [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]

                    stats = {'paragraph': 0, 'table_cell': 0, 'global_replace': 0,
                             'header_replace': 0, 'unknown': 0, 'failed': 0}
                    modifications_count = 0
                    buffer = ''
                    try:
                        async for chunk in llm.astream(messages):
                            if not chunk:
                                continue
                            token = chunk.content if hasattr(chunk, 'content') and isinstance(chunk.content, str) else str(chunk)
                            if not token:
                                continue
                            buffer += token
                            while '\n' in buffer:
                                line, buffer = buffer.split('\n', 1)
                                line = line.strip()
                                if not line or line == '===END===':
                                    buffer = ''
                                    break
                                mod = cx.parse_ndjson_line(line)
                                if mod is None:
                                    continue
                                modifications_count += 1
                                progress = base_progress + int((modifications_count / 30) * (end_progress - base_progress) * 0.8)
                                try:
                                    mod_type = mod.get('type', '')
                                    reason = (mod.get('reason', '') or '')[:60]
                                    if mod_type == 'paragraph':
                                        idx = int(mod.get('index', -1))
                                        new_text = mod.get('new_text', '')
                                        if cx.apply_paragraph_replace(doc, idx, new_text):
                                            stats['paragraph'] += 1
                                            yield await send({"type": "modification", "mod_type": "paragraph",
                                                "dept": dept, "location": f"[{dept}] 段落 P{idx}",
                                                "reason": reason, "preview": (new_text or '')[:60],
                                                "progress": min(progress, end_progress - 2), "index": modifications_count})
                                        else:
                                            stats['failed'] += 1
                                    elif mod_type == 'table_cell':
                                        t_idx = int(mod.get('table', -1)); r_idx = int(mod.get('row', -1)); c_idx = int(mod.get('col', -1))
                                        new_text = mod.get('new_text', '')
                                        if cx.apply_table_cell_replace(doc, t_idx, r_idx, c_idx, new_text):
                                            stats['table_cell'] += 1
                                            yield await send({"type": "modification", "mod_type": "table_cell",
                                                "dept": dept, "location": f"[{dept}] 表格 T{t_idx} 第{r_idx+1}行第{c_idx+1}列",
                                                "reason": reason, "preview": (new_text or '')[:60],
                                                "progress": min(progress, end_progress - 2), "index": modifications_count})
                                        else:
                                            stats['failed'] += 1
                                    elif mod_type == 'global_replace':
                                        old = mod.get('old', ''); new = mod.get('new', '')
                                        n = cx.apply_global_replace(doc, old, new)
                                        stats['global_replace'] += n
                                        yield await send({"type": "modification", "mod_type": "global_replace",
                                            "dept": dept, "location": f"[{dept}] 全文",
                                            "reason": reason, "preview": f"'{old}' → '{new}'（{n} 处）",
                                            "progress": min(progress, end_progress - 2), "index": modifications_count})
                                    elif mod_type == 'header_replace':
                                        old = mod.get('old', ''); new = mod.get('new', '')
                                        n = cx.apply_header_replace(doc, old, new)
                                        stats['header_replace'] += n
                                        yield await send({"type": "modification", "mod_type": "header_replace",
                                            "dept": dept, "location": f"[{dept}] 页眉/页脚",
                                            "reason": reason, "preview": f"'{old}' → '{new}'（{n} 处）",
                                            "progress": min(progress, end_progress - 2), "index": modifications_count})
                                    else:
                                        stats['unknown'] += 1
                                except Exception as e:
                                    stats['failed'] += 1
                    except Exception as stream_err:
                        err_msg = str(stream_err)
                        if 'timeout' in err_msg.lower():
                            yield await send({"type": "progress", "step": f"超时 {dept}",
                                "message": f"⚠️ {dept} AI调用超时，跳过", "progress": end_progress})
                            failed_files.append({"dept": dept, "filename": filename_tmpl, "reason": "AI超时"})
                            continue
                        logger.warning(f"[CXskill] {dept} LLM异常: {stream_err}")

                    if modifications_count == 0:
                        yield await send({"type": "progress", "step": f"跳过 {dept}",
                            "message": f"⚠️ {dept} AI未生成修改方案，跳过", "progress": end_progress})
                        failed_files.append({"dept": dept, "filename": filename_tmpl, "reason": "AI无方案"})
                        continue

                    # 保存文件 — 文件名用程序名称，去掉编号前缀
                    safe_dept = _re.sub(r'[\\/:*?"<>|]', '_', dept)
                    # 从原始模板文件名提取程序名称
                    # 如 "1-AAA-GM-QP-01 质量方针与组织功能以及权责管理程序-A0.doc" → "质量方针与组织功能以及权责管理程序"
                    # 如 "AAA-GA-QP-03 记录管理程序.doc" → "记录管理程序"
                    tmpl_base = os.path.splitext(filename_tmpl)[0]  # 去扩展名
                    # 去掉前面的编号前缀：数字开头 + 多段 XXX-XXX-XXX-数字 + 空格
                    tmpl_clean = _re.sub(r'^\d*-?[A-Z]+-\w+-\w+-?\d*\s*', '', tmpl_base)
                    # 也去掉纯 AAA-XX-XX-数字 格式
                    tmpl_clean = _re.sub(r'^[A-Z]+-\w+-\w+-?\d*\s*', '', tmpl_clean)
                    # 去掉末尾的版本号（如 "-A0"、"-A1"）
                    tmpl_clean = _re.sub(r'-A\d+$', '', tmpl_clean)
                    # 去掉"二级文件"前缀
                    tmpl_clean = tmpl_clean.replace('二级文件', '')
                    # 去掉 "(1)" 之类的后缀
                    tmpl_clean = _re.sub(r'\(\d+\)', '', tmpl_clean)
                    tmpl_clean = tmpl_clean.strip()
                    if not tmpl_clean:
                        tmpl_clean = '程序文件'
                    safe_tmpl_name = _re.sub(r'[\\/:*?"<>|]', '_', tmpl_clean)
                    out_filename = f"{safe_tmpl_name}_{safe_name}_{today_str}.docx"
                    output_path = os.path.join(export_dir, out_filename)
                    # 删除 LibreOffice 转换时自动添加的偶数页页眉页脚（避免空白页）
                    cx.remove_even_page_headers_footers(doc)
                    await asyncio.to_thread(doc.save, output_path)
                    # 清理 doc 转换产生的临时目录
                    if tmpl.get('need_convert') and converted:
                        try:
                            import shutil
                            shutil.rmtree(os.path.dirname(converted), ignore_errors=True)
                        except Exception:
                            pass
                    logger.info(f"[CXskill] 程序文件已生成: {output_path}")

                    download_url = f"/api/v1/documents/export-download/{out_filename}?sid={session_id_for_export}"
                    template_source_text = f"基于【{source_label}】模板生成\n部门：{dept}\n模板：{filename_tmpl}\n修改：{modifications_count} 处"

                    generated_files.append({
                        "filename": out_filename,
                        "download_url": download_url,
                        "dept": dept,
                        "display_name": tmpl_clean,  # 程序文件名称（如"质量方针与组织功能以及权责管理程序"）
                        "modifications_count": modifications_count,
                        "stats": stats,
                        "template_source_text": template_source_text
                    })

                    yield await send({
                        "type": "file_complete",
                        "dept": dept,
                        "display_name": tmpl_clean,
                        "filename": out_filename,
                        "download_url": download_url,
                        "modifications_count": modifications_count,
                        "template_source": source_label,
                        "template_filename": filename_tmpl,
                        "progress": end_progress
                    })

                except Exception as e:
                    logger.exception(f"[CXskill] {dept} 生成异常: {e}")
                    failed_files.append({"dept": dept, "filename": filename_tmpl, "reason": str(e)[:100]})
                    yield await send({"type": "progress", "step": f"失败 {dept}",
                        "message": f"❌ {dept} 生成失败: {str(e)[:80]}", "progress": end_progress})

            # 最终汇总
            # [Bug 修复] 把对话记录保存到会话历史，否则重启后前端看不到
            # 同时把文件下载信息嵌入消息，前端加载历史时能重新渲染下载按钮
            # 注意：用户点按钮不是发消息，不保存 HumanMessage，只保存 AI 回复
            try:
                from app.memory.manager import get_session_history, flush_session
                from langchain_core.messages import AIMessage as _AIMsg
                history = get_session_history(session_id_for_export)
                # AI 消息：记录生成的文件列表 + 嵌入下载信息JSON标记
                if generated_files:
                    # [Bug 修复] 文本部分只写摘要，详细文件列表通过下载按钮展示
                    # 避免消息超 8000 字被截断导致下载按钮标记丢失
                    success_count = len(generated_files)
                    failed_count = len(failed_files)
                    if failed_count > 0:
                        ai_msg = f"已生成 {success_count} 个程序文件（{failed_count} 个失败）。点击下方按钮下载："
                    else:
                        ai_msg = f"已生成 {success_count} 个程序文件。点击下方按钮下载："
                    # 嵌入下载信息（前端检测此标记重新渲染下载按钮）
                    download_info = {
                        "type": "procedure_files",
                        "files": [{"filename": f.get("filename",""), "download_url": f.get("download_url",""),
                                    "display_name": f.get("display_name", f.get("dept","")),
                                    "modifications_count": f.get("modifications_count", 0),
                                    "dept": f.get("dept","")} for f in generated_files]
                    }
                    import json as _json_for_msg
                    # [Bug 修复] 标记放在消息开头，防止消息超长被截断时标记丢失
                    ai_msg = f"<!--DOWNLOAD_INFO:{_json_for_msg.dumps(download_info, ensure_ascii=False)}-->\n\n" + ai_msg
                else:
                    ai_msg = "生成失败，未产生任何文件。"
                history.add_message(_AIMsg(content=ai_msg))
                # 更新会话时间 + flush 到磁盘
                parts = session_id_for_export.split("_", 1)
                if len(parts) == 2:
                    update_chat_time(parts[0], session_id_for_export)
                flush_session(session_id_for_export)
                logger.info(f"[CXskill] 对话记录已保存到会话历史: {session_id_for_export}")
            except Exception as save_err:
                logger.warning(f"[CXskill] 保存对话记录失败: {save_err}")

            yield await send({
                "type": "success",
                "total_files": len(generated_files),
                "failed_count": len(failed_files),
                "files": generated_files,
                "failed_files": failed_files,
                "model_used": current_model,
                "progress": 100
            })

        except Exception as e:
            logger.exception(f"[CXskill] 生成程序文件异常: {e}")
            yield await send({"type": "error", "message": f"生成失败: {str(e)}"})

    return StreamingResponse(
        sse_event_stream(), media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"}
    )


# ===== 体系调研数据存储（服务端，跨浏览器同步）=====

@router.post("/survey/save", summary="保存体系调研数据到服务器")
async def survey_save(request: Request, username: str = Depends(require_auth)):
    """保存体系调研数据到服务器（按用户隔离），实现跨浏览器同步"""
    body = await request.json()
    survey_data = body.get("survey_data", {})
    if not survey_data:
        raise HTTPException(status_code=400, detail="未提供调研数据")
    # 保存到 data/survey/{username}.json
    survey_dir = os.path.join(settings.DATA_DIR, "survey")
    os.makedirs(survey_dir, exist_ok=True)
    survey_file = os.path.join(survey_dir, f"{username}.json")
    import json as _json
    with open(survey_file, 'w', encoding='utf-8') as f:
        _json.dump(survey_data, f, ensure_ascii=False, indent=2)
    logger.info(f"[调研保存] 用户={username}, 字段数={len(survey_data)}")
    return {"success": True, "message": "调研数据已保存到服务器"}


@router.get("/survey/load", summary="加载体系调研数据")
async def survey_load(username: str = Depends(require_auth)):
    """从服务器加载体系调研数据（按用户隔离）"""
    survey_file = os.path.join(settings.DATA_DIR, "survey", f"{username}.json")
    if not os.path.exists(survey_file):
        return {"success": True, "survey_data": None, "message": "暂无保存的调研数据"}
    try:
        import json as _json
        with open(survey_file, 'r', encoding='utf-8') as f:
            survey_data = _json.load(f)
        logger.info(f"[调研加载] 用户={username}, 字段数={len(survey_data) if survey_data else 0}")
        return {"success": True, "survey_data": survey_data}
    except Exception as e:
        logger.warning(f"[调研加载] 读取失败: {e}")
        return {"success": True, "survey_data": None, "message": f"读取失败: {e}"}


@router.delete("/survey/clear", summary="清除体系调研数据")
async def survey_clear(username: str = Depends(require_auth)):
    """清除服务器上的体系调研数据"""
    survey_file = os.path.join(settings.DATA_DIR, "survey", f"{username}.json")
    if os.path.exists(survey_file):
        os.remove(survey_file)
        logger.info(f"[调研清除] 用户={username}")
    return {"success": True, "message": "调研数据已清除"}


# ===== 体系调研文档上传与AI提取 =====

@router.post("/survey/upload", summary="上传体系调研文档到临时目录")
async def survey_upload(file: UploadFile = File(...), username: str = Depends(require_auth)):
    """上传质量手册等文档到临时目录（不入知识库），用于AI提取企业信息"""
    import uuid
    allowed_ext = {".pdf", ".txt", ".docx", ".doc"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail=f"不支持的格式: {ext}，仅支持 PDF/TXT/DOCX/DOC")
    temp_dir = os.path.join(settings.DATA_DIR, "temp", "survey")
    os.makedirs(temp_dir, exist_ok=True)
    file_id = uuid.uuid4().hex[:8]
    safe_name = file.filename.replace('/', '_').replace('\\', '_')
    file_path = os.path.join(temp_dir, f"{file_id}_{safe_name}")
    with open(file_path, "wb") as f:
        content_bytes = await file.read()
        f.write(content_bytes)
    logger.info(f"[调研上传] 用户={username}, 文件={file.filename}, 路径={file_path}")
    return {"success": True, "file_path": file_path, "filename": file.filename}


@router.post("/survey/extract", summary="AI提取文档中的企业信息+识别文件类型")
async def survey_extract(request: Request, username: str = Depends(require_auth)):
    """调用当前选择的LLM模型，从上传的文档中提取企业信息，自动填充到体系调研表单
    同时识别文件类型（手册/程序文件/三层次文件/记录表格/其他），用于自动归类到知识库
    [Bug 4 修复] file_path 白名单校验，防止任意文件读取"""
    import json as _json
    import re
    body = await request.json()
    file_path = body.get("file_path", "")
    filename = body.get("filename", "")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=400, detail="文件不存在，请重新上传")
    
    # [Bug 4 修复] 白名单校验：file_path 必须在 data/temp/survey/ 目录下
    survey_temp_dir = os.path.realpath(os.path.join(settings.DATA_DIR, "temp", "survey"))
    real_file_path = os.path.realpath(file_path)
    if not real_file_path.startswith(survey_temp_dir + os.sep) and real_file_path != survey_temp_dir:
        logger.warning(f"[调研提取] 非法 file_path: {file_path} (用户={username})")
        raise HTTPException(status_code=403, detail="非法的文件路径，仅允许读取调研上传的临时文件")
    
    # [Bug 4 修复] filename 白名单校验
    if filename and not re.match(r'^[A-Za-z0-9_\u4e00-\u9fa5.\-() ]+$', filename):
        raise HTTPException(status_code=400, detail="非法的文件名")
    
    try:
        from app.rag.document import load_document
        docs = await asyncio.to_thread(load_document, file_path)
        if not docs:
            raise HTTPException(status_code=500, detail="无法读取文档内容")
        doc_text = "\n".join([d.page_content for d in docs[:5]])
        if len(doc_text) > 8000:
            doc_text = doc_text[:8000] + "\n...(文档内容已截断)"
        logger.info(f"[调研提取] 文档={filename}, 文本长度={len(doc_text)}")
        extract_prompt = '你是一个企业信息提取助手。请从以下文档内容中提取企业体系调研所需的信息，并识别文件类型。\n\n请严格按照以下JSON格式输出，只输出JSON，不要有任何其他文字。如果某个字段在文档中找不到，对应的值设为空字符串。\n\n{"file_type":"文件类型，必须是以下之一：手册/程序文件/三层次文件/记录表格/其他","file_type_name":"文件类型的具体名称，如：质量手册、文件控制程序、作业指导书等","file_dept":"文件所属部门，必须是以下之一：总经办二级/技术部二级/生产部二级/综合管理部二级/设备部二级/财务部二级/质量部二级/采购部二级/其他。根据文件内容中的部门标识（如文件编号中的GM=总经办、RD=技术部、MF=生产部、GA=综合管理部、ED=设备部、FC=财务部、QA=质量部、PU=采购部）或文件内容判断属于哪个部门。如果无法判断则填其他。","company_name":"公司全称","certs":["ISO9001"],"cert_other":"其他证书","chairman":"董事长","legal_rep":"法人代表","gm":"总经理","deputy_gm":"副总经理","mgmt_rep":"管理者代表","leader_group_leader":"贯标组长","leader_group_members":"组员","iso_office_head":"贯标办主任","iso_office_members":"成员","auditors":"内审员","products":"体系覆盖产品","process_flow":"生产流程","location":"地理位置","area":"占地面积","building_area":"建筑面积","staff_total":"正式员工人数","staff_mgmt":"管理技术人员","staff_edu":"中专以上人数","equipment":"设备情况","customers":"主要客户","address":"公司地址","contact":"联系人","phone":"电话","fax":"传真","mobile":"手机","purpose":"公司宗旨","quality_policy":"质量方针","quality_goal":"质量目标","design_dev":"有无设计开发","org":{"综合管理":{"dept":"部门","head":"负责人"},"研发技术":{"dept":"部门","head":"负责人"},"采购":{"dept":"部门","head":"负责人"},"市场":{"dept":"部门","head":"负责人"},"财务":{"dept":"部门","head":"负责人"},"制造生产":{"dept":"部门","head":"负责人"},"质量":{"dept":"部门","head":"负责人"}}}\n\n文件名：' + filename + '\n\n文档内容：\n'
        from app.agent.core import create_llm
        from langchain_core.messages import HumanMessage, SystemMessage
        llm = create_llm(short_response=True)
        messages = [
            SystemMessage(content="你是企业信息提取助手，只输出JSON格式结果，不要输出其他任何文字。"),
            HumanMessage(content=extract_prompt + doc_text)
        ]
        response = await llm.ainvoke(messages)
        ai_text = response.content.strip()
        json_match = re.search(r'\{[\s\S]*\}', ai_text)
        if not json_match:
            raise HTTPException(status_code=500, detail="AI返回格式异常，无法解析")
        json_str = json_match.group(0)
        json_str = json_str.replace('```json', '').replace('```', '').strip()
        try:
            fields = _json.loads(json_str)
        except _json.JSONDecodeError as e:
            logger.error(f"[调研提取] JSON解析失败: {e}")
            raise HTTPException(status_code=500, detail="AI提取的信息格式异常")
        # 提取文件类型信息
        file_type = fields.get('file_type', '')
        file_type_name = fields.get('file_type_name', '')
        file_dept = fields.get('file_dept', '')
        logger.info(f"[调研提取] 成功提取 {len(fields)} 个字段, 文件类型={file_type}/{file_type_name}, 部门={file_dept}")
        try:
            os.remove(file_path)
            logger.info(f"[调研提取] 已删除临时文件: {file_path}")
        except:
            pass
        return {"success": True, "fields": fields, "file_type": file_type, "file_type_name": file_type_name, "file_dept": file_dept}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"[调研提取] 异常: {e}")
        try:
            if file_path and os.path.exists(file_path):
                os.remove(file_path)
        except:
            pass
        raise HTTPException(status_code=500, detail=f"提取失败: {str(e)}")


# ===== 智能体知识库管理接口 =====



@router.delete("/agents/{agent_id}/knowledge", summary="删除智能体的知识库")

async def delete_agent_knowledge(agent_id: str, admin: str = Depends(require_admin)):

    """

    删除智能体对应的整个 ChromaDB collection（仅管理员可操作）

    在删除智能体时调用，确保知识库数据同步清理

    """

    if not agent_id:

        raise HTTPException(status_code=400, detail="agent_id 不能为空")

    result = delete_agent_collection(agent_id)

    if result["status"] == "error":

        raise HTTPException(status_code=500, detail=result["message"])

    return {"status": "success", "detail": result}





# ===== 诊断接口 =====



@router.get("/debug/collections", summary="列出所有 ChromaDB collection")

async def debug_collections(admin: str = Depends(require_admin)):

    """诊断接口：列出所有 ChromaDB collection 及其文档数
    [Bug 5 修复] 仅管理员可访问"""

    collections = list_all_collections()

    return {"collections": collections}





@router.post("/reindex", summary="重建知识库索引（切换embedding模型后使用）")

async def reindex_knowledge(agent_id: str = Query(None, description="智能体ID，为空时重建全局知识库"), admin: str = Depends(require_admin)):

    """

    重建指定知识库的所有文档索引。
    [Bug 5 修复] 仅管理员可触发

    

    切换embedding模型后（如从智谱embedding-3切换到本地bge-large-zh-v1.5），

    旧的向量数据维度不同，必须重建索引才能正常使用向量搜索。

    

    此接口会：

    1. 记录旧collection中的文档列表

    2. 删除旧collection

    3. 用新的embedding模型重新索引所有文档

    """

    result = await asyncio.to_thread(reindex_all_documents, agent_id=agent_id)

    if result["status"] == "error":

        raise HTTPException(status_code=500, detail=result["message"])

    return {"status": "success", "detail": result}





@router.get("/migrate/cleanup-collections", summary="清理异常的 ChromaDB collection")

async def cleanup_collections(admin: str = Depends(require_admin)):

    """

    清理空 collection 或有双重前缀的 collection
    [Bug 5 修复] 仅管理员可触发

    例如：agent_agent_xxx → 应该是 agent_xxx

    """

    import chromadb

    client = chromadb.PersistentClient(path=settings.CHROMA_DIR)

    collections = client.list_collections()

    cleaned = []



    for c in collections:

        name = c.name

        # 修复双重前缀：agent_agent_xxx → agent_xxx

        if name.startswith("agent_agent_"):

            correct_name = name.replace("agent_agent_", "agent_", 1)

            try:

                # 获取旧 collection 的数据

                old_data = c.get(include=["documents", "metadatas", "embeddings"])

                if old_data.get("ids"):

                    # 创建正确名称的 collection 并迁移数据

                    from app.rag.document import get_vector_store

                    # 从 agent_agent_xxx 提取真正的 agent_id
                    # [Bug 6 修复] 不能用 name.replace("agent_", "", 1)，那样得到 "agent_xxx"
                    # 再传给 get_vector_store 会拼成 "agent_agent_xxx"（又回到原名）
                    # 导致 add_documents 写入的就是即将被 delete 的旧 collection，数据全丢
                    # 正确做法：直接去掉双重前缀 "agent_agent_"，得到真正的 agent_id
                    real_agent_id = name[len("agent_agent_"):]  # 去掉 "agent_agent_" 前缀

                    new_vs = get_vector_store(agent_id=real_agent_id)

                    # 迁移文档

                    from langchain_core.documents import Document

                    docs = []

                    for i, doc_id in enumerate(old_data["ids"]):

                        doc = Document(

                            page_content=old_data["documents"][i] or "",

                            metadata=old_data["metadatas"][i] or {},

                        )

                        docs.append(doc)

                    if docs:

                        new_vs.add_documents(docs)

                    cleaned.append({"old": name, "new": correct_name, "migrated_docs": len(docs)})

                # 删除旧 collection

                client.delete_collection(name)

            except Exception as e:

                cleaned.append({"old": name, "error": str(e)})

        # 清理空 collection（除了 langchain）

        elif name != "langchain":

            try:

                count = c.count()

                if count == 0:

                    client.delete_collection(name)

                    cleaned.append({"deleted_empty": name})

            except:

                pass



    # 清理 vector_store 缓存

    from app.rag.document import reset_vector_store

    reset_vector_store()



    return {"status": "success", "cleaned": cleaned}

