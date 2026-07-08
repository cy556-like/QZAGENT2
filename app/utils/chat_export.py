"""
对话导出工具模块 (Chat Export Utilities)
=========================================

提供将对话消息导出为 Word (.docx) / PDF 的能力。

核心改进（修复"导出文件中残留 Markdown 标记、可读性极差"的问题）：
  - 内置轻量级 Markdown 解析器，将消息内容拆分为结构化块（标题、表格、列表、
    代码块、引用块、普通段落）。
  - Word 导出：表格 → 原生 Word 表格（带边框、表头底色、智能列宽）；
    标题 → Heading 样式；列表 → List Bullet / List Number；代码块 → 等宽
    字体 + 浅灰底色；行内 **粗体** / *斜体* / `代码` → 富文本 Run。
  - PDF 导出：复用同样的解析器，将表格渲染为 reportlab Table，避免直接
    输出 |---|---| 这样的纯文本。

所有导出函数均返回 bytes，可直接作为 HTTP Response 体返回。
"""

from __future__ import annotations

import re
import logging
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple

logger = logging.getLogger(__name__)


# =============================================================================
# 1. 轻量级 Markdown 解析器
# =============================================================================

# 表格分隔行：|:---:|---|:---|
_TABLE_SEP_RE = re.compile(r'^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$')
# 标题：# ~ ######
_HEADING_RE = re.compile(r'^(#{1,6})\s+(.+?)\s*#*\s*$')
# 有序列表：1. / 1) / 1、
_OL_RE = re.compile(r'^(\d+)[.)、]\s+(.+)$')
# 无序列表：- / * / + 后跟空格
_UL_RE = re.compile(r'^([-*+])\s+(.+)$')
# 引用块：>
_QUOTE_RE = re.compile(r'^>\s?(.*)$')
# 分隔线：--- / *** / ___（≥3 个）
_HR_RE = re.compile(r'^([-*_])\1{2,}\s*$')
# 中文序号标题：一、 二、 / 第X章 / 第X节 / （一）（二）
_CN_H2_RE = re.compile(r'^[一二三四五六七八九十]+、\s*.+')
_CN_H3_RE = re.compile(r'^[（(][一二三四五六七八九十]+[）)]\s*.+')
_CN_CHAPTER_RE = re.compile(r'^第[一二三四五六七八九十\d]+[章节部篇]\s*.+')


class Block:
    """解析后的内容块基类"""
    type: str = 'block'

    def __init__(self, **kw):
        self.__dict__.update(kw)


class HeadingBlock(Block):
    type = 'heading'

    def __init__(self, level: int, text: str):
        self.level = level
        self.text = text


class ParagraphBlock(Block):
    type = 'paragraph'

    def __init__(self, text: str):
        self.text = text


class ListItemBlock(Block):
    type = 'list_item'

    def __init__(self, ordered: bool, text: str, index: int = 0, marker: str = ''):
        self.ordered = ordered
        self.text = text
        self.index = index
        self.marker = marker


class TableBlock(Block):
    type = 'table'

    def __init__(self, header: List[str], rows: List[List[str]]):
        self.header = header
        self.rows = rows


class CodeBlock(Block):
    type = 'code'

    def __init__(self, code: str, lang: str = ''):
        self.code = code
        self.lang = lang


class QuoteBlock(Block):
    type = 'quote'

    def __init__(self, text: str):
        self.text = text


class HrBlock(Block):
    type = 'hr'


def parse_markdown(text: str) -> List[Block]:
    """将一段 Markdown 文本解析为 Block 列表。

    支持的元素：标题、表格、有序/无序列表、代码块、引用、分隔线、普通段落。
    不支持的元素（HTML、链接定义等）按普通段落处理。
    """
    lines = text.replace('\r\n', '\n').replace('\r', '\n').split('\n')
    blocks: List[Block] = []

    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        stripped = raw.strip()

        # 1) 空行跳过
        if not stripped:
            i += 1
            continue

        # 2) 代码块 ```lang ... ```
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines: List[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < n:  # 跳过结束的 ```
                i += 1
            blocks.append(CodeBlock('\n'.join(code_lines), lang=lang))
            continue

        # 3) 表格（首行 | ... |，第二行 |---|---|）
        if '|' in stripped and stripped.startswith('|'):
            if i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1].strip()):
                table_lines: List[str] = []
                while i < n:
                    line = lines[i].strip()
                    # 标准表格行：以 | 开头且中间还有 |
                    if line.startswith('|') and '|' in line[1:]:
                        table_lines.append(line)
                        i += 1
                    # 容错1: 折行续行——不以 | 开头但以 | 结尾，且上一行是表格行
                    elif table_lines and line.endswith('|') and '|' in line:
                        # 把续行内容追加到上一行末尾（去掉首尾空格和尾部的|）
                        cont = line.rstrip()
                        if cont.endswith('|'):
                            cont = cont[:-1].strip()
                        if cont:
                            table_lines[-1] = table_lines[-1].rstrip('|').rstrip() + ' ' + cont + ' |'
                        i += 1
                    # 容错2: 纯续行——不以 | 开头也不以 | 结尾，但上一行表格行的单元格数不够
                    elif table_lines and not line.startswith(('#', '```', '>')) and not _HR_RE.match(line) and not _UL_RE.match(line) and not _OL_RE.match(line):
                        # 检查上一行是否单元格数不足（可能是折行的内容）
                        prev = table_lines[-1]
                        prev_cells = [c for c in prev.strip('|').split('|')]
                        # 如果上一行只有1个单元格（可能是 | xxx | 这种残缺行），把当前行作为第二个单元格
                        if len(prev_cells) <= 2 and '|' in prev:
                            table_lines[-1] = prev.rstrip('|').rstrip() + ' ' + line + ' |'
                            i += 1
                        else:
                            break
                    else:
                        break
                header, rows = _parse_table(table_lines)
                if header:
                    blocks.append(TableBlock(header, rows))
                continue

        # 4) 标题 # ... ######
        m = _HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            blocks.append(HeadingBlock(level, text))
            i += 1
            continue

        # 5) 中文序号标题
        if _CN_CHAPTER_RE.match(stripped):
            blocks.append(HeadingBlock(2, stripped))
            i += 1
            continue
        if _CN_H2_RE.match(stripped):
            blocks.append(HeadingBlock(2, stripped))
            i += 1
            continue
        if _CN_H3_RE.match(stripped):
            blocks.append(HeadingBlock(3, stripped))
            i += 1
            continue

        # 6) 分隔线
        if _HR_RE.match(stripped):
            blocks.append(HrBlock())
            i += 1
            continue

        # 7) 引用块 > ...
        if stripped.startswith('>'):
            quote_lines: List[str] = []
            while i < n:
                line = lines[i].strip()
                if line.startswith('>'):
                    quote_lines.append(_QUOTE_RE.sub(r'\1', line))
                    i += 1
                else:
                    break
            blocks.append(QuoteBlock('\n'.join(quote_lines).strip()))
            continue

        # 8) 无序列表 - / * / +
        m = _UL_RE.match(stripped)
        if m:
            while i < n:
                line = lines[i].strip()
                lm = _UL_RE.match(line)
                if lm:
                    blocks.append(ListItemBlock(False, lm.group(2).strip(), marker=lm.group(1)))
                    i += 1
                elif line and not line.startswith(('- ', '* ', '+ ')) and not _OL_RE.match(line):
                    # 列表项的续行（缩进或同段延续）
                    if line[0] in ('-', '*', '+'):
                        break
                    # 续行追加到上一个列表项
                    if blocks and isinstance(blocks[-1], ListItemBlock):
                        blocks[-1].text += ' ' + line
                        i += 1
                        continue
                    else:
                        break
                else:
                    break
            continue

        # 9) 有序列表 1. / 1)
        m = _OL_RE.match(stripped)
        if m:
            idx = 1
            while i < n:
                line = lines[i].strip()
                lm = _OL_RE.match(line)
                if lm:
                    try:
                        idx = int(lm.group(1))
                    except ValueError:
                        idx += 1
                    blocks.append(ListItemBlock(True, lm.group(2).strip(), index=idx))
                    idx += 1
                    i += 1
                elif line and not _UL_RE.match(line) and not line.startswith(('- ', '* ', '+ ')):
                    if blocks and isinstance(blocks[-1], ListItemBlock):
                        blocks[-1].text += ' ' + line
                        i += 1
                        continue
                    else:
                        break
                else:
                    break
            continue

        # 10) 普通段落：合并连续非空行
        para_lines = [stripped]
        i += 1
        while i < n:
            line = lines[i]
            s = line.strip()
            if not s:
                break
            # 遇到特殊块起始则停止
            if (s.startswith('#') or s.startswith('|') or s.startswith('```')
                    or s.startswith('>') or _HR_RE.match(s)
                    or _UL_RE.match(s) or _OL_RE.match(s)
                    or _CN_CHAPTER_RE.match(s) or _CN_H2_RE.match(s) or _CN_H3_RE.match(s)):
                break
            para_lines.append(s)
            i += 1
        blocks.append(ParagraphBlock(' '.join(para_lines)))

    return blocks


def _parse_table(table_lines: List[str]) -> Tuple[List[str], List[List[str]]]:
    """解析 Markdown 表格行 → (header, rows)

    容错处理：
    - 跳过分隔行 |---|---|
    - 单元格数不齐时按表头列数补空字符串
    """
    def split_row(line: str) -> List[str]:
        s = line.strip()
        if s.startswith('|'):
            s = s[1:]
        if s.endswith('|'):
            s = s[:-1]
        return [c.strip() for c in s.split('|')]

    if len(table_lines) < 2:
        return [], []

    header = split_row(table_lines[0])
    num_cols = len(header)
    rows: List[List[str]] = []
    for line in table_lines[2:]:  # 跳过分隔行
        if _TABLE_SEP_RE.match(line.strip()):
            continue
        row = split_row(line)
        # 补齐或截断到表头列数
        if len(row) < num_cols:
            row = row + [''] * (num_cols - len(row))
        elif len(row) > num_cols:
            # 合并多余的单元格（可能是内容里含 | 导致误切）
            row = row[:num_cols - 1] + [' | '.join(row[num_cols - 1:])]
        rows.append(row)
    return header, rows


# =============================================================================
# 2. 行内 Markdown 渲染（粗体 / 斜体 / 行内代码 / 链接）
# =============================================================================

# 匹配顺序很重要：先 `code`，再 **bold**，再 *italic*，再 [text](url)
_INLINE_TOKEN_RE = re.compile(
    r'(?P<code>`[^`]+?`)'
    r'|(?P<bold>\*\*[^*]+?\*\*)'
    r'|(?P<italic>\*[^*]+?\*)'
    r'|(?P<link>\[[^\]]+?\]\([^)]+?\))'
)


def split_inline(text: str) -> List[Tuple[str, str]]:
    """将一段含行内 Markdown 标记的文本拆分为 [(kind, content), ...]

    kind ∈ {'text', 'bold', 'italic', 'code', 'link'}
    link 的 content 为 (label, url) 元组；其它为字符串。
    """
    tokens: List[Tuple[str, str]] = []
    pos = 0
    for m in _INLINE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            tokens.append(('text', text[pos:m.start()]))
        if m.group('code'):
            tokens.append(('code', m.group('code')[1:-1]))
        elif m.group('bold'):
            tokens.append(('bold', m.group('bold')[2:-2]))
        elif m.group('italic'):
            tokens.append(('italic', m.group('italic')[1:-1]))
        elif m.group('link'):
            inner = m.group('link')
            label_m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', inner)
            if label_m:
                tokens.append(('link', (label_m.group(1), label_m.group(2))))
            else:
                tokens.append(('text', inner))
        pos = m.end()
    if pos < len(text):
        tokens.append(('text', text[pos:]))
    return tokens


def strip_markdown_inline(text: str) -> str:
    """去掉所有行内 Markdown 标记，返回纯文本（用于代码块、表头等场景）"""
    def _repl(m: re.Match) -> str:
        if m.group('code'):
            return m.group('code')[1:-1]
        if m.group('bold'):
            return m.group('bold')[2:-2]
        if m.group('italic'):
            return m.group('italic')[1:-1]
        if m.group('link'):
            inner = m.group('link')
            label_m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', inner)
            return label_m.group(1) if label_m else inner
        return m.group(0)

    return _INLINE_TOKEN_RE.sub(_repl, text)


# =============================================================================
# 3. Word (.docx) 导出
# =============================================================================

def generate_chat_docx_bytes(
    messages: List[Dict[str, Any]],
    session_id: str,
    agent_name: str = '',
) -> bytes:
    """生成对话导出 Word 文档（bytes）

    Args:
        messages: [{"role": "user"/"assistant", "content": "..."}]
        session_id: 会话 ID
        agent_name: 智能体名称（用于标题）

    Returns:
        .docx 文件的二进制内容
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ===== 页面 & 默认字体 =====
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15

    # 标题样式
    for level in range(1, 5):
        hname = f'Heading {level}'
        if hname in doc.styles:
            hs = doc.styles[hname]
            hs.font.name = '宋体'
            hs.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            hs.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
            hs.paragraph_format.space_after = Pt(4)

    # 文档主标题
    display_title = f"{agent_name} 对话记录" if agent_name else "东风科技研发智能体 对话记录"
    h = doc.add_heading(display_title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f"Session: {session_id[:12] if session_id else ''}")
    info_run.font.size = Pt(9)
    info_run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()

    USER_COLOR = RGBColor(33, 33, 33)
    ASST_COLOR = RGBColor(25, 118, 210)
    CODE_BG = 'F2F2F2'
    TABLE_HEADER_BG = 'D9E2F3'
    QUOTE_BG = 'FFF8E1'

    def _set_run_font(run, size_pt=11, bold=False, italic=False, color=None, mono=False):
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if mono:
            run.font.name = 'Consolas'
        else:
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _add_inline_runs(p, text: str, size_pt=11, mono=False):
        """将含行内 Markdown 的文本渲染为 Run 序列"""
        for kind, content in split_inline(text):
            if kind == 'text':
                run = p.add_run(content)
                _set_run_font(run, size_pt, mono=mono)
            elif kind == 'bold':
                run = p.add_run(content)
                _set_run_font(run, size_pt, bold=True, mono=mono)
            elif kind == 'italic':
                run = p.add_run(content)
                _set_run_font(run, size_pt, italic=True, mono=mono)
            elif kind == 'code':
                run = p.add_run(content)
                _set_run_font(run, size_pt - 1, mono=True, color=RGBColor(0x9C, 0x27, 0xB0))
            elif kind == 'link':
                label, url = content
                run = p.add_run(label)
                _set_run_font(run, size_pt, color=RGBColor(0x19, 0x76, 0xD2))
                run.font.underline = True

    def _shade(cell, fill: str):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def _add_table(header: List[str], rows: List[List[str]]):
        num_cols = max(len(header), max((len(r) for r in rows), default=0), 1)
        # 补齐
        header = header + [''] * (num_cols - len(header))
        rows = [r + [''] * (num_cols - len(r)) for r in rows]
        num_rows = len(rows) + 1

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 智能列宽（按内容长度加权）
        col_max_lens = [0] * num_cols
        for c_idx, c in enumerate(header):
            col_max_lens[c_idx] = max(col_max_lens[c_idx], len(c))
        for r in rows:
            for c_idx, c in enumerate(r):
                col_max_lens[c_idx] = max(col_max_lens[c_idx], len(c))
        total_chars = max(sum(col_max_lens), 1)
        # 16cm ≈ 9072 dxa
        col_widths = [max(int(9072 * (l / total_chars)), 567) for l in col_max_lens]
        # 归一化到 9072
        scale = 9072 / sum(col_widths)
        col_widths = [int(w * scale) for w in col_widths]

        # 应用列宽
        tbl = table._tbl
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            for gc in tblGrid.findall(qn('w:gridCol')):
                tblGrid.remove(gc)
            for w in col_widths:
                gc = OxmlElement('w:gridCol')
                gc.set(qn('w:w'), str(w))
                tblGrid.append(gc)

        # 表头
        for c_idx, cell_text in enumerate(header):
            cell = table.cell(0, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _shade(cell, TABLE_HEADER_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(strip_markdown_inline(cell_text))
            _set_run_font(run, 10, bold=True)

        # 数据行
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, cell_text in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                _add_inline_runs(p, cell_text, size_pt=10)

        # 表格后空段
        sp = doc.add_paragraph('')
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(2)

    def _add_code_block(code: str):
        # 单元格 1x1 模拟代码块（带浅灰背景）
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        _shade(cell, CODE_BG)
        # 清空默认空段
        cell.text = ''
        for line_idx, line in enumerate(code.split('\n')):
            if line_idx == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(line if line else ' ')
            _set_run_font(run, 9, mono=True)
        sp = doc.add_paragraph('')
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(2)

    def _add_quote_block(text: str):
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        _shade(cell, QUOTE_BG)
        cell.text = ''
        for idx, line in enumerate(text.split('\n')):
            if idx == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            _add_inline_runs(p, line, size_pt=11)
            for run in p.runs:
                run.italic = True
        sp = doc.add_paragraph('')
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(2)

    def _render_blocks(blocks: List[Block]):
        for blk in blocks:
            try:
                if isinstance(blk, HeadingBlock):
                    level = min(blk.level, 4)
                    h = doc.add_heading(strip_markdown_inline(blk.text), level=level)
                elif isinstance(blk, ParagraphBlock):
                    p = doc.add_paragraph()
                    _add_inline_runs(p, blk.text, size_pt=11)
                elif isinstance(blk, ListItemBlock):
                    style_name = 'List Number' if blk.ordered else 'List Bullet'
                    try:
                        p = doc.add_paragraph(style=style_name)
                    except KeyError:
                        p = doc.add_paragraph()
                        p.add_run(('• ' if not blk.ordered else f'{blk.index}. '))
                    _add_inline_runs(p, strip_markdown_inline(blk.text), size_pt=11)
                elif isinstance(blk, TableBlock):
                    _add_table(blk.header, blk.rows)
                elif isinstance(blk, CodeBlock):
                    _add_code_block(blk.code)
                elif isinstance(blk, QuoteBlock):
                    _add_quote_block(blk.text)
                elif isinstance(blk, HrBlock):
                    p = doc.add_paragraph('─' * 40)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning(f'[导出 Word] 渲染 block {blk.type} 失败，降级为纯文本: {e}')
                try:
                    fallback_text = getattr(blk, 'text', None) or getattr(blk, 'code', None) or ''
                    if isinstance(blk, TableBlock):
                        fallback_text = ' | '.join(blk.header) + '\n' + '\n'.join(' | '.join(r) for r in blk.rows)
                    if fallback_text:
                        p = doc.add_paragraph(str(fallback_text))
                        for run in p.runs:
                            _set_run_font(run, 10)
                except Exception:
                    pass

    # ===== 逐条消息渲染 =====
    for msg in messages:
        role = msg.get('role', 'assistant')
        content = msg.get('content', '') or ''
        role_label = '用户' if role == 'user' else '助手'

        # 角色标签
        rp = doc.add_paragraph()
        rr = rp.add_run(f'{role_label}：')
        rr.bold = True
        rr.font.size = Pt(11)
        rr.font.color.rgb = USER_COLOR if role == 'user' else ASST_COLOR
        rr.font.name = '宋体'
        rr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 解析 markdown 并渲染
        try:
            blocks = parse_markdown(content)
            if not blocks:
                # 纯文本兜底
                p = doc.add_paragraph(content)
                for run in p.runs:
                    _set_run_font(run, 10)
            else:
                _render_blocks(blocks)
        except Exception as e:
            logger.warning(f'[导出 Word] 解析消息失败，降级为纯文本: {e}')
            p = doc.add_paragraph(content)
            for run in p.runs:
                _set_run_font(run, 10)

        # 分隔线
        sep = doc.add_paragraph('─' * 40)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# =============================================================================
# 5. PDF 导出（复用 markdown 解析，表格用 reportlab Table 渲染）
# =============================================================================

def generate_chat_pdf_bytes(
    messages: List[Dict[str, Any]],
    session_id: str,
    agent_name: str = '',
) -> bytes:
    """生成对话导出 PDF（bytes）—— 解析 Markdown 后渲染为原生 PDF 元素

    重要：不再静默回退到旧版 generate_chat_pdf（旧版会把 Markdown 原文当纯文本输出）。
    如果渲染失败，直接抛出异常，让上游返回 500 错误，便于定位问题。
    """
    return _generate_pdf_reportlab(messages, session_id, agent_name)


def _generate_pdf_reportlab(messages, session_id, agent_name) -> bytes:
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
                                     Table as RLTable, TableStyle)
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    # 寻找中文字体
    from app.utils.pdf_generator import find_chinese_font, _strip_emoji
    font_path = find_chinese_font()
    if not font_path:
        raise RuntimeError('未找到中文字体')
    font_name = 'ChineseFont'
    try:
        pdfmetrics.registerFont(TTFont(font_name, font_path))
    except Exception:
        # 已注册则忽略
        pass

    title_style = ParagraphStyle('Title', fontName=font_name, fontSize=16, leading=22,
                                  alignment=TA_CENTER, spaceAfter=4 * mm)
    info_style = ParagraphStyle('Info', fontName=font_name, fontSize=9, leading=13,
                                 alignment=TA_CENTER, textColor=HexColor('#888888'), spaceAfter=8 * mm)
    role_user_style = ParagraphStyle('RoleUser', fontName=font_name, fontSize=11, leading=16,
                                      textColor=HexColor('#1a1a1a'), spaceAfter=2 * mm)
    role_asst_style = ParagraphStyle('RoleAsst', fontName=font_name, fontSize=11, leading=16,
                                      textColor=HexColor('#1976D2'), spaceAfter=2 * mm)
    body_style = ParagraphStyle('Body', fontName=font_name, fontSize=10, leading=14,
                                 spaceAfter=2 * mm, wordWrap='CJK')
    heading_style_2 = ParagraphStyle('H2', fontName=font_name, fontSize=14, leading=20,
                                      textColor=HexColor('#1976D2'), spaceBefore=4 * mm, spaceAfter=2 * mm,
                                      wordWrap='CJK')
    heading_style_3 = ParagraphStyle('H3', fontName=font_name, fontSize=12, leading=18,
                                      textColor=HexColor('#1976D2'), spaceBefore=3 * mm, spaceAfter=1.5 * mm,
                                      wordWrap='CJK')
    quote_style = ParagraphStyle('Quote', fontName=font_name, fontSize=10, leading=14,
                                  leftIndent=6 * mm, textColor=HexColor('#666666'),
                                  spaceAfter=2 * mm, wordWrap='CJK')

    def _esc(s: str) -> str:
        return (s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))

    def _inline_to_html(text: str) -> str:
        """将行内 Markdown 转 reportlab Paragraph 支持的简单 HTML"""
        out = []
        for kind, content in split_inline(text):
            if kind == 'text':
                out.append(_esc(content))
            elif kind == 'bold':
                out.append(f'<b>{_esc(content)}</b>')
            elif kind == 'italic':
                out.append(f'<i>{_esc(content)}</i>')
            elif kind == 'code':
                out.append(f'<font face="Courier" color="#9C27B0">{_esc(content)}</font>')
            elif kind == 'link':
                label, _url = content
                out.append(f'<font color="#1976D2"><u>{_esc(label)}</u></font>')
        return ''.join(out)

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=15 * mm, bottomMargin=15 * mm)
    story = []

    display_title = f"{agent_name} 对话记录" if agent_name else "东风科技研发智能体 对话记录"
    story.append(Paragraph(_esc(display_title), title_style))
    story.append(Paragraph(f"Session: {_esc(session_id[:12] if session_id else '')}", info_style))

    for msg in messages:
        role = msg.get('role', 'assistant')
        content = msg.get('content', '') or ''
        role_label = '用户' if role == 'user' else '助手'
        content = _strip_emoji(content)

        role_style = role_user_style if role == 'user' else role_asst_style
        story.append(Paragraph(f'<b>{_esc(role_label)}：</b>', role_style))

        try:
            blocks = parse_markdown(content)
        except Exception as e:
            logger.warning(f'[导出 PDF] parse_markdown 失败，降级为纯文本: {e}')
            blocks = [ParagraphBlock(content)]

        for blk in blocks:
            try:
                if isinstance(blk, HeadingBlock):
                    hs = heading_style_2 if blk.level <= 2 else heading_style_3
                    story.append(Paragraph(_inline_to_html(strip_markdown_inline(blk.text)), hs))
                elif isinstance(blk, ParagraphBlock):
                    story.append(Paragraph(_inline_to_html(blk.text), body_style))
                elif isinstance(blk, ListItemBlock):
                    prefix = f'{blk.index}. ' if blk.ordered else '• '
                    story.append(Paragraph(f'{prefix}{_inline_to_html(strip_markdown_inline(blk.text))}',
                                            ParagraphStyle('Li', parent=body_style, leftIndent=6 * mm)))
                elif isinstance(blk, TableBlock):
                    header = blk.header
                    rows = blk.rows
                    if not header:
                        continue
                    num_cols = max(len(header), max((len(r) for r in rows), default=0))
                    if num_cols == 0:
                        continue
                    header = header + [''] * (num_cols - len(header))
                    rows = [r + [''] * (num_cols - len(r)) for r in rows]

                    # 列数多时缩小字号，保证列宽够用
                    if num_cols <= 4:
                        cell_font_size = 9
                    elif num_cols <= 6:
                        cell_font_size = 8
                    elif num_cols <= 8:
                        cell_font_size = 7
                    else:
                        cell_font_size = 6
                    cell_style = ParagraphStyle('Cell', fontName=font_name,
                                                 fontSize=cell_font_size,
                                                 leading=cell_font_size + 2,
                                                 wordWrap='CJK')
                    header_cell_style = ParagraphStyle('HCell', fontName=font_name,
                                                        fontSize=cell_font_size,
                                                        leading=cell_font_size + 2,
                                                        wordWrap='CJK')

                    # 显式计算列宽：A4 纵向可用宽度 = 210 - 15*2 = 180mm
                    # 减去左右 padding（各 3pt = ~1mm），按内容长度加权分配
                    avail_width_pt = A4[0] - 30 * mm  # ~510pt
                    # 按各列内容最大字符数加权
                    col_max_lens = [0] * num_cols
                    for c_idx, c in enumerate(header):
                        col_max_lens[c_idx] = max(col_max_lens[c_idx], len(c))
                    for r in rows:
                        for c_idx, c in enumerate(r):
                            col_max_lens[c_idx] = max(col_max_lens[c_idx], len(c))
                    total_chars = max(sum(col_max_lens), 1)
                    # 每列最小宽度 = 2 个汉字 = ~24pt；最大不超过 1/3 总宽
                    min_col_pt = 24
                    max_col_pt = avail_width_pt / 3
                    raw_widths = [max(min_col_pt, min(max_col_pt,
                                                       avail_width_pt * (l / total_chars)))
                                  for l in col_max_lens]
                    # 归一化到总可用宽度
                    raw_sum = sum(raw_widths)
                    col_widths = [w * avail_width_pt / raw_sum for w in raw_widths]

                    data = [[Paragraph(_inline_to_html(strip_markdown_inline(c)) or '&nbsp;', header_cell_style) for c in header]]
                    for r in rows:
                        data.append([Paragraph(_inline_to_html(strip_markdown_inline(c)) or '&nbsp;', cell_style) for c in r])
                    tbl = RLTable(data, repeatRows=1, colWidths=col_widths)
                    tbl.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#D9E2F3')),
                        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#999999')),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTSIZE', (0, 0), (-1, -1), cell_font_size),
                        ('TOPPADDING', (0, 0), (-1, -1), 2),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                        ('LEFTPADDING', (0, 0), (-1, -1), 3),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                    ]))
                    story.append(tbl)
                    story.append(Spacer(1, 2 * mm))
                elif isinstance(blk, CodeBlock):
                    # 用 Paragraph + <br/> 代替 Preformatted，避免 XML 解析问题
                    code_html = _esc(blk.code).replace('\n', '<br/>')
                    code_para_style = ParagraphStyle('CodeP', fontName='Courier', fontSize=9, leading=12,
                                                      leftIndent=4 * mm, spaceAfter=2 * mm,
                                                      backColor=HexColor('#F2F2F2'),
                                                      borderColor=HexColor('#CCCCCC'), borderWidth=0.5,
                                                      borderPadding=4)
                    story.append(Paragraph(code_html, code_para_style))
                    story.append(Spacer(1, 2 * mm))
                elif isinstance(blk, QuoteBlock):
                    story.append(Paragraph(f'<i>{_inline_to_html(strip_markdown_inline(blk.text))}</i>', quote_style))
                elif isinstance(blk, HrBlock):
                    story.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#cccccc')))
            except Exception as e:
                logger.warning(f'[导出 PDF] 渲染 block {blk.type} 失败，降级为纯文本: {e}')
                # 降级：把 block 的文本内容当普通段落渲染
                try:
                    fallback_text = getattr(blk, 'text', None) or getattr(blk, 'code', None) or ''
                    if isinstance(blk, TableBlock):
                        fallback_text = ' | '.join(blk.header) + ' | ' + ' | '.join(' | '.join(r) for r in blk.rows)
                    if fallback_text:
                        story.append(Paragraph(_esc(str(fallback_text)), body_style))
                except Exception:
                    pass

        story.append(Spacer(1, 2 * mm))
        story.append(HRFlowable(width='100%', thickness=0.5, color=HexColor('#cccccc')))
        story.append(Spacer(1, 2 * mm))

    doc.build(story)
    return buf.getvalue()
