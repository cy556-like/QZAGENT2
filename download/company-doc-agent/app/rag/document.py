"""
文档处理与向量化模块 (RAG)
负责：加载文档 → 分块 → 向量化 → 存入 ChromaDB → 检索
优化：单例缓存 Embeddings 和 VectorStore 实例，避免重复创建
支持：PDF、TXT、DOCX、XLSX

功能清单：
- load_document: 加载文档（PDF/TXT/DOCX/XLSX）
- read_document_content: 读取文档纯文本内容
- split_documents: 文档分块
- index_document: 文档索引流程（含 agent_id 隔离）
- search_documents: 向量语义检索（含 agent_id 隔离）
- list_indexed_documents: 列出已索引文档（含 agent_id 隔离）
- get_document_content: 获取指定文档完整内容（按 agent_id 隔离）
- delete_document: 删除文档+向量分块+原始文件
- update_document: 修改文档内容并重新索引
- export_document_as_docx: 导出为 DOCX 文件
- export_document_as_xlsx: 导出为 XLSX 文件
- _search_disk_files: 磁盘文件搜索（向量搜索失败的兜底方案）
"""
import os
import logging
import time
import shutil
from typing import Optional

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings

logger = logging.getLogger(__name__)

try:
    from langchain_chroma import Chroma
except ImportError:
    from langchain_community.vectorstores import Chroma

from app.config import settings


# ===== 单例缓存 =====
_vector_store = None
_embeddings = None


def get_embeddings():
    """获取 Embedding 模型（单例缓存，避免重复创建）"""
    global _embeddings
    if _embeddings is None:
        embedding_model = getattr(settings, 'EMBEDDING_MODEL', 'embedding-3')
        _embeddings = OpenAIEmbeddings(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL,
            model=embedding_model,
        )
    return _embeddings


def get_vector_store():
    """获取 ChromaDB 向量数据库实例（单例缓存，避免重复创建）"""
    global _vector_store
    if _vector_store is None:
        embeddings = get_embeddings()
        _vector_store = Chroma(
            persist_directory=settings.CHROMA_DIR,
            embedding_function=embeddings,
        )
    return _vector_store


def load_document(file_path: str) -> list:
    """
    根据文件类型加载文档
    支持：PDF、TXT、DOCX、XLSX
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif ext == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    elif ext == ".docx":
        loader = Docx2txtLoader(file_path)
        return loader.load()
    elif ext == ".xlsx":
        return _load_xlsx(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 PDF/TXT/DOCX/XLSX")


def _load_xlsx(file_path: str) -> list:
    """
    加载 XLSX 文件为 LangChain Document 列表
    每个 Sheet 作为一个 Document
    """
    try:
        from app.utils.xlsx_handler import read_xlsx_to_text
    except ImportError:
        raise ImportError("XLSX 支持需要 openpyxl，请运行: pip install openpyxl")

    filename = os.path.basename(file_path)
    text = read_xlsx_to_text(file_path)

    # 将整个 XLSX 内容作为一个 Document
    docs = [Document(page_content=text, metadata={"source": file_path, "source_file": filename})]
    return docs


def read_document_content(file_path: str) -> str:
    """
    读取文档的纯文本内容（用于文档修改功能）

    Args:
        file_path: 文档路径

    Returns:
        str: 文档纯文本内容
    """
    docs = load_document(file_path)
    content = "\n\n".join([doc.page_content for doc in docs])
    return content


def split_documents(docs: list, chunk_size: int = 500, chunk_overlap: int = 100) -> list:
    """
    文档分块
    - chunk_size: 每块最大字符数
    - chunk_overlap: 块间重叠字符数（保证上下文连续性）
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""],
    )
    return splitter.split_documents(docs)


def _get_agent_docs_dir(agent_id: str = None) -> str:
    """获取指定 agent 的文档目录"""
    if agent_id:
        agent_dir = os.path.join(settings.DOCUMENTS_DIR, agent_id)
    else:
        agent_dir = settings.DOCUMENTS_DIR
    os.makedirs(agent_dir, exist_ok=True)
    return agent_dir


def _find_file_on_disk(filename: str, agent_id: str = None) -> Optional[str]:
    """在 agent 的文档目录中查找文件"""
    agent_dir = _get_agent_docs_dir(agent_id)
    file_path = os.path.join(agent_dir, filename)
    if os.path.exists(file_path):
        return file_path
    # 也检查全局文档目录
    global_path = os.path.join(settings.DOCUMENTS_DIR, filename)
    if os.path.exists(global_path):
        return global_path
    return None


def index_document(file_path: str, filename: str = None, agent_id: str = None) -> dict:
    """
    完整的文档索引流程：加载 → 分块 → 向量化 → 存储

    Args:
        file_path: 文档文件路径
        filename: 文档名称（默认从路径中提取）
        agent_id: 智能体ID，用于知识库隔离

    Returns:
        dict: 包含分块数量和状态信息
    """
    if filename is None:
        filename = os.path.basename(file_path)

    # 1. 加载文档
    docs = load_document(file_path)

    # 2. 给文档添加元数据（含 agent_id 隔离标记）
    for doc in docs:
        doc.metadata["source_file"] = filename
        if agent_id:
            doc.metadata["agent_id"] = agent_id

    # 3. 分块
    chunks = split_documents(docs)

    # 4. 向量化并存储
    vector_store = get_vector_store()
    vector_store.add_documents(chunks)

    return {
        "filename": filename,
        "chunks": len(chunks),
        "status": "success",
        "message": f"文档 {filename} 已成功索引，共 {len(chunks)} 个分块",
    }


def search_documents(query: str, top_k: int = 3, agent_id: str = None) -> list[dict]:
    """
    在向量数据库中检索与查询最相关的文档片段

    Args:
        query: 用户查询
        top_k: 返回最相关的 K 个结果
        agent_id: 智能体ID，用于知识库隔离

    Returns:
        list[dict]: 检索结果列表
    """
    vector_store = get_vector_store()

    # 按 agent_id 过滤检索结果
    filter_dict = {}
    if agent_id:
        filter_dict = {"agent_id": agent_id}

    try:
        if filter_dict:
            results = vector_store.similarity_search_with_score(
                query, k=top_k, filter=filter_dict
            )
        else:
            results = vector_store.similarity_search_with_score(query, k=top_k)
    except Exception as e:
        # ChromaDB filter 可能不支持某些版本，降级为无 filter
        logger.warning(f"带 filter 搜索失败，降级为无 filter: {e}")
        results = vector_store.similarity_search_with_score(query, k=top_k)

    formatted = []
    for doc, score in results:
        # 如果指定了 agent_id，二次验证结果属于当前 agent
        if agent_id and doc.metadata.get("agent_id") and doc.metadata.get("agent_id") != agent_id:
            continue
        formatted.append({
            "content": doc.page_content,
            "source": doc.metadata.get("source_file", "未知来源"),
            "relevance_score": round(1 - score, 4),  # 转换为相似度
        })

    return formatted


def list_indexed_documents(agent_id: str = None) -> list[str]:
    """
    列出知识库中所有已索引的文档

    Args:
        agent_id: 智能体ID，用于知识库隔离
    """
    vector_store = get_vector_store()
    try:
        collection = vector_store._collection
        all_docs = collection.get(include=["metadatas"])
        sources = set()
        for meta in all_docs["metadatas"]:
            if meta and "source_file" in meta:
                # 如果指定了 agent_id，只返回该 agent 的文档
                if agent_id:
                    doc_agent = meta.get("agent_id", "")
                    # 兼容：没有 agent_id 标记的旧文档也算
                    if doc_agent and doc_agent != agent_id:
                        continue
                sources.add(meta["source_file"])
        return sorted(list(sources))
    except Exception:
        return []


def get_document_content(filename: str, agent_id: str = None) -> dict:
    """
    获取指定文档的完整内容（按 agent_id 隔离）

    Args:
        filename: 文档文件名（含扩展名）
        agent_id: 智能体ID

    Returns:
        dict: 包含 status, content, char_count 等信息
    """
    # 在磁盘上查找文件
    file_path = _find_file_on_disk(filename, agent_id)

    if not file_path:
        return {"status": "not_found", "message": f"文档 {filename} 未找到"}

    try:
        content = read_document_content(file_path)
        if not content or not content.strip():
            return {"status": "empty", "message": f"文档 {filename} 内容为空"}
        return {
            "status": "success",
            "content": content,
            "char_count": len(content),
            "file_path": file_path,
        }
    except Exception as e:
        return {"status": "error", "message": f"读取文档失败: {str(e)}"}


def delete_document(filename: str, agent_id: str = None) -> dict:
    """
    从知识库中删除指定文档，同时移除其所有向量分块和原始文件

    Args:
        filename: 文档文件名（含扩展名）
        agent_id: 智能体ID

    Returns:
        dict: 删除结果
    """
    # 1. 从向量数据库中删除该文档的所有分块
    vector_store = get_vector_store()
    try:
        collection = vector_store._collection
        # 查找属于该文档的所有分块 ID
        filter_dict = {"source_file": filename}
        if agent_id:
            filter_dict["agent_id"] = agent_id

        try:
            results = collection.get(where=filter_dict, include=["metadatas"])
        except Exception:
            # ChromaDB filter 不支持时，降级获取
            results = collection.get(where={"source_file": filename}, include=["metadatas"])

        chunk_ids = results.get("ids", [])
        if not chunk_ids:
            return {"status": "not_found", "message": f"文档 {filename} 在知识库中未找到"}

        # 删除向量分块
        collection.delete(ids=chunk_ids)
        logger.info(f"已从向量库删除 {len(chunk_ids)} 个分块: {filename}")

    except Exception as e:
        logger.error(f"删除向量分块失败: {e}")
        # 继续尝试删除文件

    # 2. 删除原始文件
    file_path = _find_file_on_disk(filename, agent_id)
    if file_path and os.path.exists(file_path):
        try:
            os.remove(file_path)
            logger.info(f"已删除原始文件: {file_path}")
        except Exception as e:
            logger.error(f"删除原始文件失败: {e}")

    return {
        "status": "success",
        "message": f"文档 {filename} 已从知识库和磁盘删除（共移除 {len(chunk_ids) if chunk_ids else 0} 个向量分块）",
    }


def update_document(filename: str, content: str, agent_id: str = None, async_reindex: bool = True) -> dict:
    """
    修改文档内容并重新索引到向量数据库

    Args:
        filename: 文档文件名（含扩展名）
        content: 新的文档内容
        agent_id: 智能体ID
        async_reindex: 是否异步重新索引（默认 True）

    Returns:
        dict: 修改结果
    """
    # 1. 查找原始文件
    file_path = _find_file_on_disk(filename, agent_id)
    if not file_path:
        return {"status": "not_found", "message": f"文档 {filename} 未找到"}

    # 2. 写入新内容
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"已更新文档内容: {file_path} ({len(content)} 字符)")
    except Exception as e:
        return {"status": "error", "message": f"写入文档失败: {str(e)}"}

    # 3. 删除旧的向量分块
    vector_store = get_vector_store()
    try:
        collection = vector_store._collection
        filter_dict = {"source_file": filename}
        if agent_id:
            filter_dict["agent_id"] = agent_id

        try:
            results = collection.get(where=filter_dict)
        except Exception:
            results = collection.get(where={"source_file": filename})

        old_ids = results.get("ids", [])
        if old_ids:
            collection.delete(ids=old_ids)
            logger.info(f"已删除旧向量分块 {len(old_ids)} 个")
    except Exception as e:
        logger.warning(f"删除旧向量分块失败（继续重新索引）: {e}")

    # 4. 重新索引
    try:
        result = index_document(file_path, filename=filename, agent_id=agent_id)
        logger.info(f"文档重新索引完成: {filename}, {result.get('chunks', 0)} 个分块")
    except Exception as e:
        logger.error(f"重新索引失败: {e}")
        return {"status": "error", "message": f"文档已更新但重新索引失败: {str(e)}"}

    return {
        "status": "success",
        "message": f"文档 {filename} 已更新并重新索引（{result.get('chunks', 0)} 个分块）",
    }


def export_document_as_docx(content: str, filename: str, title: str = "", session_id: str = None) -> dict:
    """
    将文本内容导出为 DOCX 文件

    Args:
        content: 文档内容（Markdown格式）
        filename: 输出文件名
        title: 文档标题
        session_id: 会话ID，用于按会话隔离存储

    Returns:
        dict: 包含 status, filename, file_path
    """
    try:
        # 导出目录：static/exports/{session_id}/
        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        if session_id:
            export_dir = os.path.join(static_dir, "exports", session_id)
        else:
            export_dir = os.path.join(static_dir, "exports")
        os.makedirs(export_dir, exist_ok=True)

        if not filename.endswith('.docx'):
            filename += '.docx'

        output_path = os.path.join(export_dir, filename)

        # 尝试使用 python-docx 生成专业格式
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import re

            doc = Document()

            # 设置默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(11)

            # 添加标题
            if title:
                heading = doc.add_heading(title, level=1)
                for run in heading.runs:
                    run.font.name = '黑体'

            # 解析 Markdown 内容并转换为 Word 格式
            lines = content.split('\n')
            i = 0
            in_table = False
            table_rows = []

            while i < len(lines):
                line = lines[i]

                # 标题识别
                if line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                    i += 1
                    continue
                elif line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                    i += 1
                    continue
                elif line.startswith('# '):
                    doc.add_heading(line[2:].strip(), level=1)
                    i += 1
                    continue

                # Markdown 表格识别
                if '|' in line and line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                        table_rows = []

                    # 跳过分隔行 |---|---|
                    if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                        i += 1
                        continue

                    # 收集表格行
                    cells = [c.strip() for c in line.strip().strip('|').split('|')]
                    table_rows.append(cells)
                    i += 1

                    # 检查下一行是否还是表格
                    if i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                        continue
                    else:
                        # 表格结束，写入 Word 表格
                        if table_rows:
                            max_cols = max(len(row) for row in table_rows)
                            table = doc.add_table(rows=len(table_rows), cols=max_cols, style='Table Grid')
                            for row_idx, row_data in enumerate(table_rows):
                                for col_idx, cell_text in enumerate(row_data):
                                    if col_idx < max_cols:
                                        cell = table.cell(row_idx, col_idx)
                                        cell.text = cell_text
                                        # 表头加粗
                                        if row_idx == 0:
                                            for paragraph in cell.paragraphs:
                                                for run in paragraph.runs:
                                                    run.bold = True
                        in_table = False
                        table_rows = []
                    continue

                # 列表项
                if line.strip().startswith('- ') or line.strip().startswith('* '):
                    text = line.strip()[2:].strip()
                    # 处理粗体 **text**
                    p = doc.add_paragraph(style='List Bullet')
                    _add_formatted_text(p, text)
                    i += 1
                    continue

                # 普通段落
                if line.strip():
                    p = doc.add_paragraph()
                    _add_formatted_text(p, line.strip())
                else:
                    # 空行
                    pass

                i += 1

            doc.save(output_path)

        except ImportError:
            # python-docx 不可用，保存为 TXT
            txt_filename = filename.rsplit('.', 1)[0] + '.txt'
            output_path = os.path.join(export_dir, txt_filename)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            filename = txt_filename

        return {
            "status": "success",
            "filename": filename,
            "file_path": output_path,
            "message": f"文档已导出为 {filename}",
        }

    except Exception as e:
        logger.error(f"导出 DOCX 失败: {e}")
        return {"status": "error", "message": f"导出失败: {str(e)}"}


def _add_formatted_text(paragraph, text: str):
    """向 Word 段落添加带格式的文本（支持 **粗体** 标记）"""
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def export_document_as_xlsx(content: str, filename: str, title: str = "", session_id: str = None) -> dict:
    """
    将文本内容导出为 XLSX 文件

    Args:
        content: 文档内容（Markdown表格格式）
        filename: 输出文件名
        title: 工作表标题
        session_id: 会话ID，用于按会话隔离存储

    Returns:
        dict: 包含 status, filename, file_path
    """
    try:
        from app.utils.xlsx_handler import write_xlsx_from_text

        # 导出目录
        static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static")
        if session_id:
            export_dir = os.path.join(static_dir, "exports", session_id)
        else:
            export_dir = os.path.join(static_dir, "exports")
        os.makedirs(export_dir, exist_ok=True)

        if not filename.endswith('.xlsx'):
            filename = filename.rsplit('.', 1)[0] + '.xlsx'

        output_path = os.path.join(export_dir, filename)

        actual_path = write_xlsx_from_text(content, output_path, sheet_name=title or "Sheet1")
        actual_filename = os.path.basename(actual_path)

        return {
            "status": "success",
            "filename": actual_filename,
            "file_path": actual_path,
            "message": f"Excel文档已导出为 {actual_filename}",
        }

    except Exception as e:
        logger.error(f"导出 XLSX 失败: {e}")
        return {"status": "error", "message": f"导出失败: {str(e)}"}


def _search_disk_files(query: str, top_k: int = 5, agent_id: str = None) -> list[dict]:
    """
    磁盘文件搜索（向量搜索失败的兜底方案）
    按关键词匹配文件名和文件内容

    Args:
        query: 搜索查询
        top_k: 返回最相关的 K 个结果
        agent_id: 智能体ID

    Returns:
        list[dict]: 检索结果列表
    """
    results = []
    agent_dir = _get_agent_docs_dir(agent_id)

    if not os.path.exists(agent_dir):
        return results

    # 简单关键词匹配
    keywords = set(query.replace("？", "").replace("的", "").replace("了", "").replace("是", "").replace("什么", ""))

    for filename in os.listdir(agent_dir):
        file_path = os.path.join(agent_dir, filename)
        if not os.path.isfile(file_path):
            continue

        ext = os.path.splitext(filename)[1].lower()
        if ext not in {".pdf", ".txt", ".docx", ".xlsx", ".md"}:
            continue

        try:
            # 读取文件内容（限制长度避免内存占用过大）
            content = read_document_content(file_path)
            content_preview = content[:2000] if content else ""

            # 计算简单匹配分数
            score = 0
            for kw in keywords:
                if kw in filename:
                    score += 2  # 文件名匹配权重更高
                if kw in content_preview:
                    score += 1

            if score > 0:
                results.append({
                    "content": content_preview,
                    "source": filename,
                    "relevance_score": min(score / max(len(keywords), 1), 1.0),
                })
        except Exception:
            continue

    # 按分数排序
    results.sort(key=lambda x: x.get("relevance_score", 0), reverse=True)
    return results[:top_k]
